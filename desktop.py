import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext
from tkinter.font import Font
import requests
import os
import re
import json
import pandas as pd
from docx import Document
import docx.shared
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import ImageReader
import datetime

# os.environ['TCL_LIBRARY'] = r'C:\Users\User\AppData\Local\Programs\Python\Python311\tcl\tcl8.6'
# os.environ['TK_LIBRARY'] = r'C:\Users\User\AppData\Local\Programs\Python\Python311\tcl\tk8.6'

class ModernStyle:
    """Стиль оформления приложения с цветовой схемой и параметрами"""
    def __init__(self):
        self.primary_color = "#1e3d6d"
        self.secondary_color = "#2d5ba9"
        self.accent_color = "#40E0D0"
        self.background_color = "#f0f2f5"
        self.surface_color = "#ffffff"
        self.text_primary = "#1a1a1a"
        self.text_secondary = "#666666"
        self.border_color = "#d1d5db"
        self.success_color = "#10b981"
        self.error_color = "#ef4444"
        self.border_radius = 8

class ModernAuthApp:
    """Основное приложение авторизации с GUI интерфейсом"""
    def __init__(self, root):
        self.root = root
        self.root.title('ANTech - Авторизация')
        self.root.geometry('1920x1080')
        self.root.state('zoomed')
        self.style = ModernStyle()
        self.setup_styles()
        self.root.configure(bg=self.style.background_color)
        self.main_frame = tk.Frame(self.root, bg=self.style.background_color)
        self.main_frame.place(relx=0.5, rely=0.5, anchor='center')
        self.create_auth_widgets()
        self.base_url = 'http://127.0.0.1:8000'
        self.token = None
    
    def setup_styles(self):
        """Настройка стилей виджетов приложения"""   
        style = ttk.Style()
        style.configure('Modern.TFrame', background=self.style.background_color)
        style.configure('Auth.TFrame', background=self.style.surface_color)
        style.configure('Primary.TButton', 
                       background=self.style.primary_color,
                       foreground='#1e3d6d',
                       borderwidth=0,
                       focuscolor='none',
                       padding=(20, 10))
        style.map('Primary.TButton',
                 background=[('active', self.style.secondary_color),
                           ('pressed', self.style.primary_color)])
        style.configure('Secondary.TButton',
                       background=self.style.surface_color,
                       foreground=self.style.primary_color,
                       borderwidth=1,
                       padding=(20, 10))
        style.map('Secondary.TButton',
                 background=[('active', self.style.background_color)])
        style.configure('Modern.TLabel',
                       background=self.style.surface_color,
                       foreground=self.style.text_primary,
                       font=('Arial', 11))
        style.configure('Title.TLabel',
                       background=self.style.surface_color,
                       foreground=self.style.text_primary,
                       font=('Arial', 24, 'bold'))
        style.configure('Subtitle.TLabel',
                       background=self.style.surface_color,
                       foreground=self.style.text_secondary,
                       font=('Arial', 14))
        style.configure('Modern.TEntry',
                       fieldbackground=self.style.surface_color,
                       foreground=self.style.text_primary,
                       borderwidth=1,
                       relief='solid',
                       padding=(10, 8))
        style.map('Modern.TEntry',
                 fieldbackground=[('focus', self.style.surface_color)])
    
    def on_code_login_success(self, token, expires_at):
        """Обработка успешного входа по коду"""
        self.token = token
        self.token_expires = expires_at
        self.open_main_app()
    
    def show_code_login(self):
        """Показывает окно входа по коду"""
        CodeLoginWindow(self.root, self.style, self.base_url, self.on_code_login_success)
    
    def create_auth_widgets(self):
        """Создает виджеты для окна авторизации"""
        auth_container = ttk.Frame(self.main_frame, style='Auth.TFrame', padding=40)
        auth_container.pack(padx=20, pady=20)
        
        ttk.Label(auth_container, text='Войти', style='Title.TLabel').pack(pady=(0, 10))
        ttk.Label(auth_container, text='Войдите в свой аккаунт', style='Subtitle.TLabel').pack(pady=(0, 30))
        
        login_frame = ttk.Frame(auth_container, style='Auth.TFrame')
        login_frame.pack(fill='x', pady=(0, 15))
        
        ttk.Label(login_frame, text='Телефон или e-mail', style='Modern.TLabel').pack(anchor='w', pady=(0, 5))
        self.login_entry = ttk.Entry(login_frame, width=30, style='Modern.TEntry', font=('Arial', 12))
        self.login_entry.pack(fill='x', ipady=8)
        
        password_frame = ttk.Frame(auth_container, style='Auth.TFrame')
        password_frame.pack(fill='x', pady=(0, 10))
        
        ttk.Label(password_frame, text='Введите пароль', style='Modern.TLabel').pack(anchor='w', pady=(0, 5))
        self.password_entry = ttk.Entry(password_frame, width=30, show='*', style='Modern.TEntry', font=('Arial', 12))
        self.password_entry.pack(fill='x', ipady=8)
        
        forgot_frame = ttk.Frame(auth_container, style='Auth.TFrame')
        forgot_frame.pack(fill='x', pady=(0, 30))
        
        forgot_label = tk.Label(forgot_frame, text='Забыли пароль?', 
                               bg=self.style.surface_color, fg=self.style.primary_color,
                               font=('Arial', 11), cursor='hand2')
        forgot_label.pack(anchor='e')
        forgot_label.bind('<Button-1>', lambda e: self.show_forgot_password())
        
        login_btn = ttk.Button(auth_container, text='Войти', 
                              command=self.login, style='Primary.TButton')
        login_btn.pack(fill='x', pady=(0, 20))
        
        code_btn = ttk.Button(auth_container, text='Войти по коду',
                             command=self.show_code_login, style='Secondary.TButton')
        code_btn.pack(fill='x')
        
        register_btn = ttk.Button(auth_container, text='Зарегистрироваться',
                                 command=self.show_registration, style='Secondary.TButton')
        register_btn.pack(fill='x', pady=(10, 0))
        
        self.login_entry.bind('<Return>', lambda e: self.login())
        self.password_entry.bind('<Return>', lambda e: self.login())
    
    def show_registration(self):
        """Показывает окно регистрации"""
        RegistrationWindow(self.root, self.style, self.base_url)
    
    def is_phone(self, text):
        """Проверяет, является ли текст телефонным номером"""
        phone_pattern = r'^[\d\s\-\+\(\)]+$'
        return bool(re.match(phone_pattern, text)) and len(text) >= 5
        
    def login(self):
        """Выполняет авторизацию пользователя"""
        login = self.login_entry.get().strip()
        password = self.password_entry.get().strip()
    
        if not login or not password:
            messagebox.showerror('Ошибка!', 'Заполните все поля.')
            return
            
        try:
            if self.is_phone(login):
                auth_data = {'phone': login, 'password': password}
            else:
                auth_data = {'email': login, 'password': password}
            
            response = requests.post(
                f'{self.base_url}/users/auth/',
                json=auth_data
            )
            
            if response.status_code == 200:
                data = response.json()
                self.token = data.get('token')
                self.token_expires = data.get('expires_at')
                self.open_main_app()
            else:
                error = response.json().get('detail', 'Ошибка авторизации')
                messagebox.showerror('Ошибка!', f'Ошибка авторизации: {error}')
        except requests.exceptions.RequestException as e:
            messagebox.showerror('Ошибка!', f'Не удалось подключиться к серверу: {e}')
    
    def open_main_app(self):
        """Открывает основное приложение после успешной авторизации"""
        self.root.destroy()
        root = tk.Tk()
        MainApp(root, self.token)
        root.mainloop()
       
    def show_forgot_password(self):
        """Показывает окно восстановления пароля"""
        ForgotPasswordWindow(self.root, self.style)

class ForgotPasswordWindow:
    """Окно восстановления пароля пользователя"""
    def __init__(self, parent, style):
        self.style = style
        self.base_url = 'http://127.0.0.1:8000'
        self.email = None
        
        self.window = tk.Toplevel(parent)
        self.window.title('Восстановление пароля')
        self.window.geometry('500x500')
        self.window.configure(bg=self.style.background_color)
        self.window.resizable(False, False)
        self.window.transient(parent)
        self.window.grab_set()
        
        self.create_email_step()
    
    def create_email_step(self):
        """Создает интерфейс для ввода email"""
        if hasattr(self, 'current_frame'):
            self.current_frame.destroy()
            
        self.current_frame = ttk.Frame(self.window, style='Auth.TFrame', padding=30)
        self.current_frame.pack(fill='both', expand=True)
        
        ttk.Label(self.current_frame, text='Восстановление пароля', style='Title.TLabel').pack(pady=(0, 10))
        ttk.Label(self.current_frame, text='Введите email для получения кода подтверждения', 
                 style='Subtitle.TLabel').pack(pady=(0, 30))
        
        email_frame = ttk.Frame(self.current_frame, style='Auth.TFrame')
        email_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Label(email_frame, text='Email', style='Modern.TLabel').pack(anchor='w', pady=(0, 5))
        self.email_entry = ttk.Entry(email_frame, style='Modern.TEntry', font=('Arial', 12))
        self.email_entry.pack(fill='x', ipady=8)
        
        ttk.Button(self.current_frame, text='Отправить код', 
                  command=self.send_code, style='Primary.TButton').pack(fill='x', pady=(0, 10))
        
        ttk.Button(self.current_frame, text='Отмена', 
                  command=self.window.destroy, style='Secondary.TButton').pack(fill='x')
        
        self.email_entry.bind('<Return>', lambda e: self.send_code())
    
    def create_code_step(self):
        """Создает интерфейс для ввода кода подтверждения"""
        self.current_frame.destroy()
        
        self.current_frame = ttk.Frame(self.window, style='Auth.TFrame', padding=30)
        self.current_frame.pack(fill='both', expand=True)
        
        ttk.Label(self.current_frame, text='Восстановление пароля', style='Title.TLabel').pack(pady=(0, 10))
        ttk.Label(self.current_frame, text=f'Код отправлен на {self.email}', 
                 style='Subtitle.TLabel').pack(pady=(0, 30))
        
        code_frame = ttk.Frame(self.current_frame, style='Auth.TFrame')
        code_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Label(code_frame, text='Код подтверждения', style='Modern.TLabel').pack(anchor='w', pady=(0, 5))
        self.code_entry = ttk.Entry(code_frame, style='Modern.TEntry', font=('Arial', 12))
        self.code_entry.pack(fill='x', ipady=8)
        
        password_frame = ttk.Frame(self.current_frame, style='Auth.TFrame')
        password_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Label(password_frame, text='Новый пароль', style='Modern.TLabel').pack(anchor='w', pady=(0, 5))
        self.new_password_entry = ttk.Entry(password_frame, show='*', style='Modern.TEntry', font=('Arial', 12))
        self.new_password_entry.pack(fill='x', ipady=8)
        
        ttk.Button(self.current_frame, text='Сменить пароль', 
                  command=self.confirm_password_change, style='Primary.TButton').pack(fill='x', pady=(0, 10))
        
        ttk.Button(self.current_frame, text='Назад', 
                  command=self.create_email_step, style='Secondary.TButton').pack(fill='x')
        
        self.code_entry.bind('<Return>', lambda e: self.confirm_password_change())
        self.new_password_entry.bind('<Return>', lambda e: self.confirm_password_change())
    
    def send_code(self):
        """Отправляет код подтверждения на email"""
        email = self.email_entry.get().strip()
        
        if not email:
            messagebox.showerror('Ошибка!', 'Введите email')
            return
            
        try:
            response = requests.post(
                f'{self.base_url}/users/change_password/',
                params={'email': email}
            )
            
            if response.status_code == 200:
                self.email = email
                self.create_code_step()
            else:
                error = response.json().get('detail', 'Ошибка отправки кода')
                messagebox.showerror('Ошибка!', error)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror('Ошибка!', f'Не удалось подключиться к серверу: {e}')
    
    def confirm_password_change(self):
        """Подтверждает смену пароля с помощью кода"""
        code = self.code_entry.get().strip()
        new_password = self.new_password_entry.get().strip()
        
        if not code or not new_password:
            messagebox.showerror('Ошибка!', 'Заполните все поля')
            return
            
        try:
            response = requests.post(
                f'{self.base_url}/users/confirm_change_password/',
                params={
                    'email': self.email,
                    'code': code,
                    'new_password': new_password
                }
            )
            
            if response.status_code == 200:
                messagebox.showinfo('Успех!', 'Пароль успешно изменен')
                self.window.destroy()
            else:
                error = response.json().get('detail', 'Ошибка смены пароля')
                messagebox.showerror('Ошибка!', error)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror('Ошибка!', f'Не удалось подключиться к серверу: {e}')

class CodeLoginWindow:
    """Окно входа по одноразовому коду"""
    def __init__(self, parent, style, base_url, on_success):
        self.style = style
        self.base_url = base_url
        self.on_success = on_success
        self.email = None
        
        self.window = tk.Toplevel(parent)
        self.window.title('Вход по коду')
        self.window.geometry('500x500')
        self.window.configure(bg=self.style.background_color)
        self.window.resizable(False, False)
        self.window.transient(parent)
        self.window.grab_set()
        
        self.create_email_step()
    
    def create_email_step(self):
        """Ввод email"""
        if hasattr(self, 'current_frame'):
            self.current_frame.destroy()
            
        self.current_frame = ttk.Frame(self.window, style='Auth.TFrame', padding=30)
        self.current_frame.pack(fill='both', expand=True)
        
        ttk.Label(self.current_frame, text='Вход по коду', style='Title.TLabel').pack(pady=(0, 10))
        ttk.Label(self.current_frame, text='Введите email для получения кода подтверждения', 
                 style='Subtitle.TLabel').pack(pady=(0, 30))
        
        email_frame = ttk.Frame(self.current_frame, style='Auth.TFrame')
        email_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Label(email_frame, text='Email', style='Modern.TLabel').pack(anchor='w', pady=(0, 5))
        self.email_entry = ttk.Entry(email_frame, style='Modern.TEntry', font=('Arial', 12))
        self.email_entry.pack(fill='x', ipady=8)
        
        ttk.Button(self.current_frame, text='Получить код', 
                  command=self.send_code, style='Primary.TButton').pack(fill='x', pady=(0, 10))
        
        ttk.Button(self.current_frame, text='Отмена', 
                  command=self.window.destroy, style='Secondary.TButton').pack(fill='x')
        
        self.email_entry.bind('<Return>', lambda e: self.send_code())
    
    def create_code_step(self):
        """Ввод кода"""
        self.current_frame.destroy()
        
        self.current_frame = ttk.Frame(self.window, style='Auth.TFrame', padding=30)
        self.current_frame.pack(fill='both', expand=True)
        
        ttk.Label(self.current_frame, text='Вход по коду', style='Title.TLabel').pack(pady=(0, 10))
        ttk.Label(self.current_frame, text=f'Код отправлен на {self.email}', 
                 style='Subtitle.TLabel').pack(pady=(0, 30))
        
        code_frame = ttk.Frame(self.current_frame, style='Auth.TFrame')
        code_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Label(code_frame, text='Код подтверждения', style='Modern.TLabel').pack(anchor='w', pady=(0, 5))
        self.code_entry = ttk.Entry(code_frame, style='Modern.TEntry', font=('Arial', 12))
        self.code_entry.pack(fill='x', ipady=8)
        
        ttk.Button(self.current_frame, text='Войти', 
                  command=self.confirm_login, style='Primary.TButton').pack(fill='x', pady=(0, 10))
        
        ttk.Button(self.current_frame, text='Назад', 
                  command=self.create_email_step, style='Secondary.TButton').pack(fill='x')
        
        self.code_entry.bind('<Return>', lambda e: self.confirm_login())
    
    def send_code(self):
        """Отправка кода на email"""
        email = self.email_entry.get().strip()
        
        if not email:
            messagebox.showerror('Ошибка!', 'Введите email')
            return
            
        try:
            response = requests.post(
                f'{self.base_url}/users/request_login_code/',
                params={'email': email}
            )
            
            if response.status_code == 200:
                self.email = email
                self.create_code_step()
            else:
                error = response.json().get('detail', 'Ошибка отправки кода')
                messagebox.showerror('Ошибка!', error)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror('Ошибка!', f'Не удалось подключиться к серверу: {e}')
    
    def confirm_login(self):
        """Подтверждение входа по коду"""
        code = self.code_entry.get().strip()
        
        if not code:
            messagebox.showerror('Ошибка!', 'Введите код подтверждения')
            return
            
        try:
            response = requests.post(
                f'{self.base_url}/users/confirm_login_code/',
                params={'email': self.email, 'code': code}
            )
            
            if response.status_code == 200:
                data = response.json()
                if self.window and self.window.winfo_exists():
                    self.window.destroy()
                self.on_success(data.get('token'), data.get('expires_at'))
            else:
                error = response.json().get('detail', 'Ошибка входа')
                messagebox.showerror('Ошибка!', error)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror('Ошибка!', f'Не удалось подключиться к серверу: {e}')

class RegistrationWindow:
    """Окно регистрации нового пользователя"""
    def __init__(self, parent, style, base_url):
        self.style = style
        self.base_url = base_url
        
        self.window = tk.Toplevel(parent)
        self.window.title('Регистрация')
        self.window.geometry('600x900')
        self.window.configure(bg=self.style.background_color)
        self.window.resizable(False, False)
        self.window.transient(parent)
        self.window.grab_set()
        
        self.create_widgets()
    
    def create_widgets(self):
        """Создает виджеты для регистрации"""
        container = ttk.Frame(self.window, style='Auth.TFrame', padding=30)
        container.pack(fill='both', expand=True)
        
        ttk.Label(container, text='Регистрация', style='Title.TLabel').pack(pady=(0, 10))
        ttk.Label(container, text='Создайте новый аккаунт', style='Subtitle.TLabel').pack(pady=(0, 30))
        
        name_frame = ttk.Frame(container, style='Auth.TFrame')
        name_frame.pack(fill='x', pady=(0, 15))
        ttk.Label(name_frame, text='ФИО', style='Modern.TLabel').pack(anchor='w', pady=(0, 5))
        self.name_entry = ttk.Entry(name_frame, style='Modern.TEntry', font=('Arial', 12))
        self.name_entry.pack(fill='x', ipady=8)
        
        email_frame = ttk.Frame(container, style='Auth.TFrame')
        email_frame.pack(fill='x', pady=(0, 15))
        ttk.Label(email_frame, text='Email', style='Modern.TLabel').pack(anchor='w', pady=(0, 5))
        self.email_entry = ttk.Entry(email_frame, style='Modern.TEntry', font=('Arial', 12))
        self.email_entry.pack(fill='x', ipady=8)
        
        phone_frame = ttk.Frame(container, style='Auth.TFrame')
        phone_frame.pack(fill='x', pady=(0, 15))
        ttk.Label(phone_frame, text='Телефон', style='Modern.TLabel').pack(anchor='w', pady=(0, 5))
        self.phone_entry = ttk.Entry(phone_frame, style='Modern.TEntry', font=('Arial', 12))
        self.phone_entry.pack(fill='x', ipady=8)
        
        password_frame = ttk.Frame(container, style='Auth.TFrame')
        password_frame.pack(fill='x', pady=(0, 15))
        ttk.Label(password_frame, text='Пароль', style='Modern.TLabel').pack(anchor='w', pady=(0, 5))
        self.password_entry = ttk.Entry(password_frame, show='*', style='Modern.TEntry', font=('Arial', 12))
        self.password_entry.pack(fill='x', ipady=8)
        
        address_frame = ttk.Frame(container, style='Auth.TFrame')
        address_frame.pack(fill='x', pady=(0, 30))
        ttk.Label(address_frame, text='Адрес', style='Modern.TLabel').pack(anchor='w', pady=(0, 5))
        self.address_entry = ttk.Entry(address_frame, style='Modern.TEntry', font=('Arial', 12))
        self.address_entry.pack(fill='x', ipady=8)
        
        register_btn = ttk.Button(container, text='Зарегистрироваться', 
                                command=self.register, style='Primary.TButton')
        register_btn.pack(fill='x', pady=(0, 20))
        
        cancel_btn = ttk.Button(container, text='Отмена', 
                              command=self.window.destroy, style='Secondary.TButton')
        cancel_btn.pack(fill='x')
        
        self.name_entry.bind('<Return>', lambda e: self.email_entry.focus())
        self.email_entry.bind('<Return>', lambda e: self.phone_entry.focus())
        self.phone_entry.bind('<Return>', lambda e: self.password_entry.focus())
        self.password_entry.bind('<Return>', lambda e: self.address_entry.focus())
        self.address_entry.bind('<Return>', lambda e: self.register())
    
    def register(self):
        """Выполняет регистрацию нового пользователя"""
        name = self.name_entry.get().strip()
        email = self.email_entry.get().strip()
        phone = self.phone_entry.get().strip()
        password = self.password_entry.get().strip()
        address = self.address_entry.get().strip()
        
        if not all([name, email, phone, password, address]):
            messagebox.showerror('Ошибка!', 'Заполните все поля')
            return
            
        try:
            response = requests.post(
                f'{self.base_url}/users/register/',
                params={
                    'email': email,
                    'password': password,
                    'full_name': name,
                    'number_phone': phone,
                    'address': address
                }
            )
            
            if response.status_code == 200:
                messagebox.showinfo('Успех!', 'Регистрация прошла успешно!')
                self.window.destroy()
            else:
                error = response.json().get('detail', 'Ошибка регистрации')
                messagebox.showerror('Ошибка!', error)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror('Ошибка!', f'Не удалось подключиться к серверу: {e}')

class MainApp:
    """Основное приложение после авторизации"""
    def __init__(self, root, token):
        self.root = root
        self.root.title('ANTech - Конфигуратор ПК')
        self.root.state('zoomed')
        self.style = ModernStyle()
        self.setup_styles()
        self.root.configure(bg=self.style.background_color)
        self.token = token
        self.user_data = None
        self.base_url = 'http://127.0.0.1:8000'
        self.load_user_data()
        
        if self.user_data and self.user_data.get('role') == 'Администратор':
            self.open_admin_app()
            return
         
        self.main_container = ttk.Frame(root, style='Modern.TFrame')
        self.main_container.pack(fill='both', expand=True)
        
        self.create_header()
        self.create_tabs()
        
        self.init_catalog_tab()
        self.init_configurations_tab()
        self.init_orders_tab()
        self.init_profile_tab()
    
    def setup_styles(self):
        """Настройка стилей виджетов"""
        style = ttk.Style()
        style.configure('Modern.TFrame', background=self.style.background_color)
        style.configure('Surface.TFrame', background=self.style.surface_color)
        style.configure('Header.TLabel', 
                       font=('Arial', 16, 'bold'),
                       background=self.style.surface_color,
                       foreground=self.style.text_primary)
        style.configure('Normal.TLabel', 
                       font=('Arial', 11),
                       background=self.style.surface_color,
                       foreground=self.style.text_primary)
        style.configure('Secondary.TLabel',
                       font=('Arial', 11),
                       background=self.style.surface_color,
                       foreground=self.style.text_secondary)
        style.configure('Primary.TButton', 
                       font=('Arial', 11), 
                       background=self.style.primary_color,
                       foreground='#1e3d6d',
                       borderwidth=0,
                       focuscolor='none')
        style.map('Primary.TButton',
                 background=[('active', self.style.secondary_color)])
        style.configure('Secondary.TButton',
                       font=('Arial', 11),
                       background=self.style.surface_color,
                       foreground=self.style.primary_color,
                       borderwidth=1)
        style.map('Secondary.TButton',
                 background=[('active', self.style.background_color)])
        style.configure('Modern.TNotebook', 
                       background=self.style.background_color)
        style.configure('Modern.TNotebook.Tab', 
                       background=self.style.surface_color,
                       foreground=self.style.text_primary,
                       padding=(20, 10))
        style.map('Modern.TNotebook.Tab',
                 background=[('selected', self.style.primary_color),
                           ('active', self.style.secondary_color)],
                 foreground=[('selected', 'white')])
        style.configure('Treeview',
                       background=self.style.surface_color,
                       foreground=self.style.text_primary,
                       fieldbackground=self.style.surface_color,
                       rowheight=25)
        style.configure('Treeview.Heading',
                       background=self.style.primary_color,
                       foreground="#1e3d6d",
                       font=('Arial', 10, 'bold'))
        style.map('Treeview.Heading',
                 background=[('active', self.style.secondary_color)])
        style.configure('Modern.TEntry',
                       fieldbackground=self.style.surface_color,
                       foreground=self.style.text_primary,
                       borderwidth=1,
                       relief='solid')
        style.configure('Modern.TCombobox',
                       fieldbackground=self.style.surface_color,
                       foreground=self.style.text_primary)
    
    def create_header(self):
        """Создает заголовок приложения"""
        header_frame = ttk.Frame(self.main_container, style='Surface.TFrame', padding=20)
        header_frame.pack(fill='x', padx=20, pady=10)
        
        title_frame = ttk.Frame(header_frame, style='Surface.TFrame')
        title_frame.pack(side='left')
        
        ttk.Label(title_frame, text='ANTech', 
                 style='Header.TLabel').pack(anchor='w')
        ttk.Label(title_frame, text=f'Добро пожаловать, {self.user_data["name"]}!', 
                 style='Secondary.TLabel').pack(anchor='w')
        
        ttk.Button(header_frame, text='Выход', 
                  command=self.logout, style='Secondary.TButton').pack(side='right')
    
    def create_tabs(self):
        """Создает вкладки интерфейса"""
        tab_control = ttk.Notebook(self.main_container, style='Modern.TNotebook')
        
        self.tab_catalog = ttk.Frame(tab_control, style='Surface.TFrame', padding=20)
        self.tab_configurations = ttk.Frame(tab_control, style='Surface.TFrame', padding=20)
        self.tab_orders = ttk.Frame(tab_control, style='Surface.TFrame', padding=20)
        self.tab_profile = ttk.Frame(tab_control, style='Surface.TFrame', padding=20)
        
        tab_control.add(self.tab_catalog, text='Каталог компонентов')
        tab_control.add(self.tab_configurations, text='Мои конфигурации')
        tab_control.add(self.tab_orders, text='Мои заказы')
        tab_control.add(self.tab_profile, text='Профиль')
        
        tab_control.pack(fill='both', expand=True, padx=20, pady=10)

    def load_user_data(self):
        """Загружает данные пользователя"""
        try:
            headers = {'token': self.token}
            response = requests.get(f'{self.base_url}/users/me/', headers=headers)
            if response.status_code == 200:
                self.user_data = response.json()
            else:
                messagebox.showerror('Ошибка!', 'Не удалось получить данные пользователя')
                self.logout()
        except requests.exceptions.RequestException as e:
            messagebox.showerror('Ошибка!', f'Ошибка соединения с сервером: {e}')
            self.logout()
    
    def logout(self):
        """Выполняет выход из системы"""
        for widget in self.root.winfo_children():
            widget.destroy()
        ModernAuthApp(self.root)
    
    def open_admin_app(self):
        """Открывает административную панель"""
        for widget in self.root.winfo_children():
            widget.destroy()
        AdminApp(self.root, self.token)
    
    def init_catalog_tab(self):
        """Инициализирует вкладку каталога компонентов"""
        ttk.Label(self.tab_catalog, text='Каталог компонентов', style='Header.TLabel').pack(anchor='w', pady=(0, 20))
        
        filter_container = ttk.Frame(self.tab_catalog, style='Surface.TFrame', padding=15)
        filter_container.pack(fill='x', pady=(0, 20))
        
        ttk.Label(filter_container, text='Фильтры', style='Normal.TLabel').pack(anchor='w', pady=(0, 10))
        
        filter_frame = ttk.Frame(filter_container, style='Surface.TFrame')
        filter_frame.pack(fill='x')
        
        type_frame = ttk.Frame(filter_frame, style='Surface.TFrame')
        type_frame.pack(side='left', padx=(0, 15))
        ttk.Label(type_frame, text='Тип:', style='Secondary.TLabel').pack(anchor='w')
        self.type_filter = tk.StringVar()
        self.type_combo = ttk.Combobox(type_frame, textvariable=self.type_filter, 
                                      state='readonly', width=15, style='Modern.TCombobox')
        self.type_combo.pack(pady=(5, 0))
        
        price_frame = ttk.Frame(filter_frame, style='Surface.TFrame')
        price_frame.pack(side='left', padx=(0, 15))
        ttk.Label(price_frame, text='Цена:', style='Secondary.TLabel').pack(anchor='w')
        
        price_subframe = ttk.Frame(price_frame, style='Surface.TFrame')
        price_subframe.pack(pady=(5, 0))
        
        self.min_price_filter = tk.StringVar()
        min_price_entry = ttk.Entry(price_subframe, textvariable=self.min_price_filter, 
                                   width=8, style='Modern.TEntry')
        min_price_entry.pack(side='left')
        ttk.Label(price_subframe, text='-', style='Secondary.TLabel').pack(side='left', padx=5)
        self.max_price_filter = tk.StringVar()
        max_price_entry = ttk.Entry(price_subframe, textvariable=self.max_price_filter, 
                                   width=8, style='Modern.TEntry')
        max_price_entry.pack(side='left')
        
        button_frame = ttk.Frame(filter_frame, style='Surface.TFrame')
        button_frame.pack(side='left', padx=(20, 0))
        
        ttk.Button(button_frame, text='Применить', 
                  command=self.apply_filters, style='Primary.TButton').pack(side='left', padx=(0, 5))
        ttk.Button(button_frame, text='Сбросить', 
                  command=self.reset_filters, style='Secondary.TButton').pack(side='left', padx=(0, 5))
        ttk.Button(button_frame, text='Экспорт',
                   command=self.export_catalog, style='Primary.TButton').pack(side='left')
        
        search_frame = ttk.Frame(self.tab_catalog, style='Surface.TFrame')
        search_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Label(search_frame, text='Поиск:', style='Secondary.TLabel').pack(anchor='w', pady=(0, 5))
        
        search_subframe = ttk.Frame(search_frame, style='Surface.TFrame')
        search_subframe.pack(fill='x')
        
        self.search_var = tk.StringVar()
        search_entry = ttk.Entry(search_subframe, textvariable=self.search_var, 
                                style='Modern.TEntry')
        search_entry.pack(side='left', fill='x', expand=True, padx=(0, 10))
        
        ttk.Button(search_subframe, text='Найти', 
                  command=self.search_components, style='Primary.TButton').pack(side='left')
        
        table_frame = ttk.Frame(self.tab_catalog, style='Surface.TFrame')
        table_frame.pack(fill='both', expand=True)
        
        columns = ('id', 'name', 'type', 'manufacture', 'price', 'stock')
        self.tree = ttk.Treeview(table_frame, columns=columns, show='headings', height=15)
        
        self.tree.heading('id', text='ID')
        self.tree.heading('name', text='Название')
        self.tree.heading('type', text='Тип')
        self.tree.heading('manufacture', text='Производитель')
        self.tree.heading('price', text='Цена')

        self.tree.heading('stock', text='Наличие')
        
        self.tree.column('id', width=50)
        self.tree.column('name', width=200)
        self.tree.column('type', width=150)
        self.tree.column('manufacture', width=150)
        self.tree.column('price', width=100)
        self.tree.column('stock', width=100)
        
        scrollbar = ttk.Scrollbar(table_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscroll=scrollbar.set)
        
        self.tree.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')
        
        self.tree.bind('<<TreeviewSelect>>', self.show_component_specification)
        
        spec_container = ttk.Frame(self.tab_catalog, style='Surface.TFrame', padding=15)
        spec_container.pack(fill='x', pady=(20, 0))
        
        ttk.Label(spec_container, text='Характеристики', style='Normal.TLabel').pack(anchor='w', pady=(0, 10))
        
        self.spec_text = scrolledtext.ScrolledText(spec_container, height=8, state='disabled',
                                                  bg=self.style.surface_color, fg=self.style.text_primary, 
                                                  font=('Arial', 10), wrap=tk.WORD)
        self.spec_text.pack(fill='x')
        
        self.all_components = []
        self.filtered_components = []
        self.available_types = set()
        
        self.load_components()

    def load_components(self):
        """Загружает список компонентов"""
        try:
            headers = {'token': self.token}
            response = requests.get(f'{self.base_url}/components/get_all/', headers=headers)
            
            if response.status_code == 200: 
                for item in self.tree.get_children():
                    self.tree.delete(item)
                
                components = response.json()
                
                self.all_components = components
                self.filtered_components = components.copy()

                self.available_types = set()
                for component in components:
                    if component.get('type_name'):
                        self.available_types.add(component['type_name'])

                self.type_combo['values'] = ["Все типы"] + sorted(list(self.available_types))
                self.type_combo.set("Все типы")

                self.display_filtered_components()
            else:
                messagebox.showerror('Ошибка!', 'Не удалось загрузить данные о компонентах')
        except requests.exceptions.RequestException as e:
            messagebox.showerror('Ошибка!', f'Не удалось получить данные от сервера: {e}')
    
    def apply_filters(self):
        """Применяет фильтры к каталогу"""
        if not hasattr(self, 'all_components') or not self.all_components:
            return
        
        filtered = self.all_components.copy()

        selected_type = self.type_filter.get()
        if selected_type and selected_type != "Все типы":
            filtered = [comp for comp in filtered if comp.get('type_name') == selected_type]

        min_price = self.min_price_filter.get()
        if min_price:
            try:
                min_price_val = float(min_price)
                filtered = [comp for comp in filtered if comp.get('price', 0) >= min_price_val]
            except ValueError:
                pass 

        max_price = self.max_price_filter.get()
        if max_price:
            try:
                max_price_val = float(max_price)
                filtered = [comp for comp in filtered if comp.get('price', 0) <= max_price_val]
            except ValueError:
                pass

        search_term = self.search_var.get().strip().lower()
        if search_term:
            filtered = [comp for comp in filtered if (
                search_term in comp['name'].lower() or 
                search_term in comp['type_name'].lower() or 
                search_term in comp['manufacture_name'].lower()
            )]
        
        self.filtered_components = filtered
        self.display_filtered_components()

    def reset_filters(self):
        """Сбрасывает фильтры каталога"""
        self.type_filter.set('')
        self.min_price_filter.set('')
        self.max_price_filter.set('')
        self.search_var.set('')
        self.filtered_components = self.all_components.copy()
        self.display_filtered_components()

    def display_filtered_components(self):
        """Отображает отфильтрованные компоненты"""
        for item in self.tree.get_children():
            self.tree.delete(item)
        
        for component in self.filtered_components:
            self.tree.insert('', 'end', values=(
                component['id'],
                component['name'],
                component['type_name'],
                component['manufacture_name'],
                component['price'],
                component['stock_quantity']
            ))
    
    def export_catalog(self):
        """Создает интерфейс для экспорта данных"""
        self.dialog = tk.Toplevel(self.root)
        self.dialog.title('Экспорт данных')
        self.dialog.geometry('400x150')
        self.dialog.configure(bg=self.style.background_color)
        self.dialog.resizable(False, False)
        self.dialog.transient(self.root)
        self.dialog.grab_set()
        
        content_frame = ttk.Frame(self.dialog, style='Surface.TFrame', padding=20)
        content_frame.pack(fill='both', expand=True)

        input_frame = ttk.Frame(content_frame, style='Surface.TFrame')
        input_frame.pack(fill='x', pady=(0, 15))
        
        ttk.Label(input_frame, text='Введите название файла:', style='Normal.TLabel').pack(side='left', padx=(0, 10))
        
        self.filename_entry = ttk.Entry(input_frame, style='Modern.TEntry', font=('Arial', 12))
        self.filename_entry.pack(fill='x', expand=True, ipady=8)

        button_frame = ttk.Frame(content_frame, style='Surface.TFrame')
        button_frame.pack(fill='x')
        
        export_btn = ttk.Button(
            button_frame, 
            text='Экспортировать', 
            style='Primary.TButton',
            command=self.perform_export 
        )
        export_btn.pack(ipady=8)
        
        
    def perform_export(self):
        """Выполняет экспорт отфильтрованного каталога компонентов в xlsx формат"""
        
        def check_filename(filename):
            """Проверка имени файла на содержание запретных символов"""
            invalid_chars = r'[<>:"/\\|?*]'
            file = re.sub(invalid_chars, '_', filename.strip())
            return file
        
        os.makedirs('./Экспорт', exist_ok=True)
        self.safe_filename = check_filename(self.filename_entry.get())
        file_path = f'./Экспорт/{self.safe_filename}.xlsx'
        if os.path.exists(file_path):
            messagebox.showwarning('Внимание!', 'Таблица с экспортированными данными уже создана')
            return
        
        try:
            df = pd.DataFrame(self.filtered_components)
            column_mapping = {
                "id": "ID",
                "name": "Название",
                "type_name": "Тип",
                "manufacture_name": "Производитель",
                "price": "Цена",
                "stock_quantity": "Наличие",
                "specification": "Характеристики"
            }
            df = df.rename(columns=column_mapping)

            with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Компоненты')

                worksheet = writer.sheets['Компоненты']
                for column in worksheet.columns:
                    max_length = 0
                    column_letter = column[0].column_letter

                    for cell in column:
                        try:
                            if len(str(cell.value)) >= max_length:
                                max_length = len(str(cell.value))
                        except:
                            pass

                    adjusted_width = (max_length + 2) * 1.2
                    worksheet.column_dimensions[column_letter].width = adjusted_width 
            
            messagebox.showinfo('Успех!', f'Каталог компонентов успешно импортирован в файл {self.safe_filename}.xlsx')
            
            self.dialog.destroy()

        except Exception as e:
            messagebox.showerror('Ошибка!', f'Ошибка при экспорте каталога компонентов: {e}')
        
    def search_components(self):
        """Выполняет поиск компонентов"""
        self.apply_filters()
        
        if self.search_var.get().strip():
            messagebox.showinfo('Поиск', f'Найдено компонентов: {len(self.filtered_components)}')


    def clear_search(self):
        """Очищает поисковый запрос"""
        self.search_var.set('')
        self.apply_filters()

    def update_components_table(self, components):
        """Обновляет таблицу компонентов"""
        for item in self.tree.get_children():
            self.tree.delete(item)

        for component in components:
            self.tree.insert('', 'end', values=(
                component['id'],
                component['name'],
                component['type_name'],
                component['manufacture_name'],
                component['price'],
                component['stock_quantity']
            ))

        self.filtered_components_data = components
    
    def show_component_specification(self, event):
        """Показывает спецификацию выбранного компонента"""
        selected_item = self.tree.selection()
        if not selected_item:
            return
        
        item = selected_item[0]
        item_index = self.tree.index(item)       

        components_to_use = self.filtered_components
        
        if item_index >= len(components_to_use):
            return
            
        component = components_to_use[item_index]
        specification = component.get('specification')
        self.spec_text.config(state='normal')
        self.spec_text.delete(1.0, tk.END)
        
        if specification:
            try:
                formated_spec = json.dumps(specification, indent=2, ensure_ascii=False)
                self.spec_text.insert(tk.END, formated_spec)
            except:
                self.spec_text.insert(tk.END, str(specification))
        else:    
            self.spec_text.insert(tk.END, 'Дополнительная информация о товаре отсутствует.')
            
        self.spec_text.config(state='disabled')
    
    def init_configurations_tab(self):
        """Инициализирует вкладку конфигураций"""
        main_frame = ttk.Frame(self.tab_configurations, style='Surface.TFrame')
        main_frame.pack(fill='both', expand=True)
        ttk.Label(main_frame, text='Мои конфигурации ПК', style='Header.TLabel').pack(anchor='w', pady=(0, 20))

        button_frame = ttk.Frame(main_frame, style='Surface.TFrame')
        button_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Button(button_frame, text='Создать новую', 
                command=self.create_new_configuration, style='Primary.TButton').pack(side='left', padx=(0, 10))
        ttk.Button(button_frame, text='Обновить', 
                command=self.load_configurations, style='Secondary.TButton').pack(side='left', padx=(0, 10))
        ttk.Button(button_frame, text='Удалить', 
                command=self.delete_configuration, style='Secondary.TButton').pack(side='left')

        paned_window = ttk.PanedWindow(main_frame, orient=tk.HORIZONTAL)
        paned_window.pack(fill='both', expand=True, pady=10)

        left_frame = ttk.Frame(paned_window, style='Surface.TFrame', padding=10)
        paned_window.add(left_frame, weight=1)
        
        ttk.Label(left_frame, text='Мои конфигурации:', style='Normal.TLabel').pack(anchor='w', pady=(0, 10))

        columns = ('id', 'name', 'description', 'created')
        self.config_tree = ttk.Treeview(left_frame, columns=columns, show='headings', height=15)
        
        self.config_tree.heading('id', text='ID')
        self.config_tree.heading('name', text='Название')
        self.config_tree.heading('description', text='Описание')
        self.config_tree.heading('created', text='Дата создания')
        
        self.config_tree.column('id', width=50)
        self.config_tree.column('name', width=150)
        self.config_tree.column('description', width=200)
        self.config_tree.column('created', width=120)
        
        scrollbar_left = ttk.Scrollbar(left_frame, orient=tk.VERTICAL, command=self.config_tree.yview)
        self.config_tree.configure(yscroll=scrollbar_left.set)
        
        self.config_tree.pack(side='left', fill='both', expand=True)
        scrollbar_left.pack(side='right', fill='y')

        right_frame = ttk.Frame(paned_window, style='Surface.TFrame', padding=10)
        paned_window.add(right_frame, weight=2)
        
        ttk.Label(right_frame, text='Сборка конфигурации:', style='Normal.TLabel').pack(anchor='w', pady=(0, 10))

        config_columns = ('type', 'component', 'price', 'quantity', 'total')
        self.config_components_tree = ttk.Treeview(right_frame, columns=config_columns, show='headings', height=10)
        
        self.config_components_tree.heading('type', text='Тип')
        self.config_components_tree.heading('component', text='Компонент')
        self.config_components_tree.heading('price', text='Цена')
        self.config_components_tree.heading('quantity', text='Кол-во')
        self.config_components_tree.heading('total', text='Итого')
        
        self.config_components_tree.column('type', width=100)
        self.config_components_tree.column('component', width=200)
        self.config_components_tree.column('price', width=80)
        self.config_components_tree.column('quantity', width=60)
        self.config_components_tree.column('total', width=80)
        
        scrollbar_right = ttk.Scrollbar(right_frame, orient=tk.VERTICAL, command=self.config_components_tree.yview)
        self.config_components_tree.configure(yscroll=scrollbar_right.set)
        
        self.config_components_tree.pack(fill='both', expand=True)
        scrollbar_right.pack(side='right', fill='y')

        bottom_frame = ttk.Frame(right_frame, style='Surface.TFrame')
        bottom_frame.pack(fill='x', pady=10)
        
        self.total_label = ttk.Label(bottom_frame, text='Общая сумма: 0 руб.', style='Header.TLabel')
        self.total_label.pack(side='left', padx=10)

        manage_frame = ttk.Frame(bottom_frame, style='Surface.TFrame')
        manage_frame.pack(side='right', padx=10)
        
        ttk.Button(manage_frame, text='Добавить компонент', 
                command=self.add_component_to_config, style='Primary.TButton').pack(side='left', padx=5)
        ttk.Button(manage_frame, text='Удалить компонент', 
                command=self.remove_component_from_config, style='Secondary.TButton').pack(side='left', padx=5)
        ttk.Button(manage_frame, text='Экспорт',
                   command=self.export_configuration_to_docx, style='Primary.TButton').pack(side='left', padx=5)
        ttk.Button(manage_frame, text='Создать заказ', 
                command=self.create_order_from_config, style='Primary.TButton').pack(side='left', padx=5)

        self.config_tree.bind('<<TreeviewSelect>>', self.on_configuration_select)
        self.config_components_tree.bind('<<TreeviewSelect>>', self.on_component_select)

        self.load_configurations()

    def load_configurations(self):
        """Загружает список конфигураций"""
        try:
            headers = {'token': self.token}
            response = requests.get(f'{self.base_url}/configurations/get_all/', headers=headers)
            
            if response.status_code == 200:
                for item in self.config_tree.get_children():
                    self.config_tree.delete(item)
                
                configurations = response.json()
                self.configurations_data = {}
                
                for config in configurations:
                    config_id = config['id']
                    self.configurations_data[config_id] = config

                    created_at = config.get('created_at', '')
                    if created_at:
                        try:
                            created_at = created_at.split('T')[0]
                        except:
                            pass
                    
                    self.config_tree.insert('', 'end', values=(
                        config_id,
                        config['name_config'] or f'Конфигурация #{config_id}',
                        config['description'] or '-',
                        created_at
                    ))
            else:
                messagebox.showerror('Ошибка!', 'Не удалось загрузить конфигурации')
        except requests.exceptions.RequestException as e:
            messagebox.showerror('Ошибка!', f'Ошибка соединения: {e}')

    def on_configuration_select(self, event):
        """Обрабатывает выбор конфигурации"""
        selected = self.config_tree.selection()
        if not selected:
            return
        
        item = selected[0]
        config_id = self.config_tree.item(item)['values'][0]
        self.current_config_id = config_id

        self.load_configuration_components(config_id)

    def load_configuration_components(self, config_id):
        """Загружает компоненты выбранной конфигурации"""
        try:
            headers = {'token': self.token}
            response = requests.get(f'{self.base_url}/configurations/{config_id}/components/', headers=headers)
            
            if response.status_code == 200:
                for item in self.config_components_tree.get_children():
                    self.config_components_tree.delete(item)
                
                components = response.json()
                total_amount = 0
                
                for component in components:
                    total_price = component['total_price']
                    total_amount += total_price
                    
                    self.config_components_tree.insert('', 'end', values=(
                        component['type_name'] or '-',
                        component['component_name'],
                        f"{component['price']:.2f} руб.",
                        component['quantity'],
                        f"{total_price:.2f} руб."
                    ), tags=(component['id'],))
                
                self.total_label.config(text=f'Общая сумма: {total_amount:.2f} руб.')
            else:
                messagebox.showerror('Ошибка!', 'Не удалось загрузить компоненты конфигурации')
        except requests.exceptions.RequestException as e:
            messagebox.showerror('Ошибка!', f'Ошибка соединения: {e}')

    def on_component_select(self, event):
        """Обрабатывает выбор компонента"""
        self.selected_component_id = None
        selected = self.config_components_tree.selection()
        if selected:
            item = selected[0]
            self.selected_component_id = self.config_components_tree.item(item)['tags'][0]

    def export_configuration_to_docx(self):
        """Экспортирует выбранную конфигурацию в формате DOCX"""
        if not hasattr(self, 'current_config_id') or not self.current_config_id:
            messagebox.showwarning('Внимание!', 'Выберите конфигурацию для экспорта')
            return

        config_data = self.configurations_data.get(self.current_config_id)
        if not config_data:
            messagebox.showerror('Ошибка!', 'Не удалось получать данные конфигурации')
            return
        
        try:
            os.makedirs('./Экспорт', exist_ok=True)
            config_name = config_data.get('name_config', f'Конфигурация_{self.current_config_id}')
            safe_filename = re.sub(r'[<>:"/\\|?*]', '_', config_name.strip())

            headers = {'token': self.token}
            response = requests.get(
                f'{self.base_url}/configurations/{self.current_config_id}/components/',
                headers=headers
            )

            if response.status_code != 200:
                messagebox.showerror('Ошибка!', 'Не удалось получить компоненты конфигурации')
                return
            
            components = response.json()

            doc = Document()

            title = doc.add_heading(f'Конфигурация: {config_name}', level=1)
            title.alignment = 1

            created_at = config_data.get('created_at', '')
            if created_at:
                try:
                    created_date = created_at.split("T")[0]
                    date_para = doc.add_paragraph(f'Дата создания: {created_date}.')
                    date_para.alignment = 1
                except:
                    pass

            doc.add_paragraph()

            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'

            header_cells = table.rows[0].cells
            header_cells[0].text = '№'
            header_cells[1].text = 'Компоненты'
            header_cells[2].text = 'Тип'
            header_cells[3].text = 'Количество'
            header_cells[4].text = 'Стоимость'

            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            total_amount = 0
            for i, component in enumerate(components, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = component.get('component_name')
                row_cells[2].text = component.get('type_name')
                row_cells[3].text = str(component.get('quantity'))
                
                component_price = component.get('total_price', 0)
                row_cells[4].text = f'{str(component_price)} руб.'

                total_amount += component_price
            
            doc.add_paragraph()

            total_para = doc.add_paragraph()
            total_para.alignment = 2
            total_run = total_para.add_run(f'Итоговая стоимость: {str(total_amount)} рублей.')
            total_run.bold = True
            total_run.font.size = docx.shared.Pt(14)

            doc.add_paragraph()

            note_para = doc.add_paragraph('Примечание: ')
            note_para.add_run('Стоимость указана на момент формирования документа и может изменяться.').italic = True

            file_path = f'./Экспорт/{safe_filename}.docx'
            doc.save(file_path)

            messagebox.showinfo('Успех!', f'Документ с данными о конфигурации успешно сформирован.\nНазвание файла: {safe_filename}.docx')
        
        except Exception as e:
            messagebox.showerror('Ошибка!', f'Ошибка при экспорте конфигурации: {e}')

    def create_new_configuration(self):
        """Создает новую конфигурацию"""
        dialog = tk.Toplevel(self.root)
        dialog.title('Создание новой конфигурации')
        dialog.geometry('500x300')
        dialog.configure(bg=self.style.background_color)
        dialog.transient(self.root)
        dialog.grab_set()
        
        content_frame = ttk.Frame(dialog, style='Surface.TFrame', padding=20)
        content_frame.pack(fill='both', expand=True)
        
        ttk.Label(content_frame, text='Название конфигурации:', style='Normal.TLabel').pack(anchor='w', pady=5)
        name_entry = ttk.Entry(content_frame, width=40, style='Modern.TEntry')
        name_entry.pack(fill='x', pady=5)
        
        ttk.Label(content_frame, text='Описание (необязательно):', style='Normal.TLabel').pack(anchor='w', pady=5)
        desc_entry = ttk.Entry(content_frame, width=40, style='Modern.TEntry')
        desc_entry.pack(fill='x', pady=5)
        
        def create_config():
            """Создает запрос на создание конфигурации"""
            name = name_entry.get().strip()
            description = desc_entry.get().strip()
            
            if not name:
                messagebox.showerror('Ошибка!', 'Введите название конфигурации')
                return
            
            try:
                headers = {'token': self.token}
                response = requests.post(
                    f'{self.base_url}/configurations/create/',
                    headers=headers,
                    json={'name_config': name, 'description': description}
                )
                
                if response.status_code == 200:
                    messagebox.showinfo('Успех!', 'Конфигурация создана')
                    dialog.destroy()
                    self.load_configurations()
                else:
                    error = response.json().get('detail', 'Ошибка создания')
                    messagebox.showerror('Ошибка!', error)
            except requests.exceptions.RequestException as e:
                messagebox.showerror('Ошибка!', f'Ошибка соединения: {e}')
        
        button_frame = ttk.Frame(content_frame, style='Surface.TFrame')
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame, text='Создать', command=create_config, style='Primary.TButton').pack(side='left', padx=10)
        ttk.Button(button_frame, text='Отмена', command=dialog.destroy, style='Secondary.TButton').pack(side='left', padx=10)

    def add_component_to_config(self):
        """Добавляет компонент в конфигурацию"""
        if not hasattr(self, 'current_config_id') or not self.current_config_id:
            messagebox.showwarning('Внимание!', 'Выберите конфигурацию для добавления компонентов')
            return

        if not hasattr(self, 'all_components') or not self.all_components:
            messagebox.showwarning('Внимание!', 'Сначала загрузите каталог компонентов')
            return
        
        dialog = tk.Toplevel(self.root)
        dialog.title('Добавление компонента')
        dialog.geometry('600x600')
        dialog.configure(bg=self.style.background_color)
        dialog.transient(self.root)
        dialog.grab_set()
        
        content_frame = ttk.Frame(dialog, style='Surface.TFrame', padding=20)
        content_frame.pack(fill='both', expand=True)
        
        ttk.Label(content_frame, text='Выберите компонент:', style='Normal.TLabel').pack(anchor='w', pady=(0, 10))

        filter_frame = ttk.Frame(content_frame, style='Surface.TFrame')
        filter_frame.pack(fill='x', pady=(0, 10))
        
        ttk.Label(filter_frame, text='Фильтр по типу:', style='Secondary.TLabel').pack(side='left')
        type_var = tk.StringVar()
        type_combo = ttk.Combobox(filter_frame, textvariable=type_var, style='Modern.TCombobox')

        types = list(set(comp['type_name'] for comp in self.all_components if comp['type_name']))
        type_combo['values'] = ['Все'] + types
        type_combo.set('Все')
        type_combo.pack(side='left', padx=5)

        columns = ('name', 'type', 'manufacture', 'price', 'stock')
        comp_tree = ttk.Treeview(content_frame, columns=columns, show='headings', height=10)
        
        comp_tree.heading('name', text='Название')
        comp_tree.heading('type', text='Тип')
        comp_tree.heading('manufacture', text='Производитель')
        comp_tree.heading('price', text='Цена')
        comp_tree.heading('stock', text='Наличие')
        
        comp_tree.column('name', width=200)
        comp_tree.column('type', width=100)
        comp_tree.column('manufacture', width=100)
        comp_tree.column('price', width=80)
        comp_tree.column('stock', width=60)
        
        scrollbar = ttk.Scrollbar(content_frame, orient=tk.VERTICAL, command=comp_tree.yview)
        comp_tree.configure(yscroll=scrollbar.set)
        
        comp_tree.pack(fill='both', expand=True, pady=5)
        scrollbar.pack(side='right', fill='y')

        quantity_frame = ttk.Frame(content_frame, style='Surface.TFrame')
        quantity_frame.pack(fill='x', pady=10)
        
        ttk.Label(quantity_frame, text='Количество:', style='Normal.TLabel').pack(side='left')
        quantity_var = tk.IntVar(value=1)
        quantity_spin = ttk.Spinbox(quantity_frame, from_=1, to=100, textvariable=quantity_var, width=10, style='Modern.TEntry')
        quantity_spin.pack(side='left', padx=5)
        
        def filter_components():
            """Фильтрация компонентов по типу"""
            for item in comp_tree.get_children():
                comp_tree.delete(item)
            
            selected_type = type_var.get()
            for component in self.all_components:
                if selected_type == 'Все' or component['type_name'] == selected_type:
                    comp_tree.insert('', 'end', values=(
                        component['name'],
                        component['type_name'] or '-',
                        component['manufacture_name'] or '-',
                        f"{component['price']:.2f} руб.",
                        component['stock_quantity']
                    ), tags=(component['name'],))
        
        def add_selected_component():
            """Добавление компонента в конфигурацию"""
            selected = comp_tree.selection()
            if not selected:
                messagebox.showwarning('Внимание!', 'Выберите компонент')
                return
            
            item = selected[0]
            component_name = comp_tree.item(item)['tags'][0]
            quantity = quantity_var.get()
            
            try:
                headers = {'token': self.token}
                response = requests.post(
                    f'{self.base_url}/configurations/{self.current_config_id}/components/',
                    headers=headers,
                    json={'component_name': component_name, 'quantity': quantity}
                )
                
                if response.status_code == 200:
                    messagebox.showinfo('Успех!', 'Компонент добавлен')
                    dialog.destroy()
                    self.load_configuration_components(self.current_config_id)
                else:
                    error = response.json().get('detail', 'Ошибка добавления')
                    messagebox.showerror('Ошибка!', error)
            except requests.exceptions.RequestException as e:
                messagebox.showerror('Ошибка!', f'Ошибка соединения: {e}')
    
        filter_components()

        type_combo.bind('<<ComboboxSelected>>', lambda e: filter_components())
        
        button_frame = ttk.Frame(content_frame, style='Surface.TFrame')
        button_frame.pack(pady=10)
        
        ttk.Button(button_frame, text='Добавить', command=add_selected_component, 
                style='Primary.TButton').pack(side='left', padx=10)
        ttk.Button(button_frame, text='Отмена', command=dialog.destroy, style='Secondary.TButton').pack(side='left', padx=10)

    def remove_component_from_config(self):
        """Удаляет компонент из конфигурации"""
        if not hasattr(self, 'current_config_id') or not self.current_config_id:
            messagebox.showwarning('Внимание!', 'Выберите конфигурацию')
            return
        
        if not hasattr(self, 'selected_component_id') or not self.selected_component_id:
            messagebox.showwarning('Внимание!', 'Выберите компонент для удаления')
            return
        
        try:
            headers = {'token': self.token}
            response = requests.delete(
                f'{self.base_url}/configurations/{self.current_config_id}/components/{self.selected_component_id}/',
                headers=headers
            )
            
            if response.status_code == 200:
                messagebox.showinfo('Успех!', 'Компонент удален')
                self.load_configuration_components(self.current_config_id)
            else:
                error = response.json().get('detail', 'Ошибка удаления')
                messagebox.showerror('Ошибка!', error)
        except requests.exceptions.RequestException as e:
            messagebox.showerror('Ошибка!', f'Ошибка соединения: {e}')

    def create_order_from_config(self):
        """Создает заказ из конфигурации"""
        if not hasattr(self, 'current_config_id') or not self.current_config_id:
            messagebox.showwarning('Внимание!', 'Выберите конфигурацию для заказа')
            return
        
        dialog = tk.Toplevel(self.root)
        dialog.title('Создание заказа')
        dialog.geometry('300x200')
        dialog.configure(bg=self.style.background_color)
        dialog.transient(self.root)
        dialog.grab_set()
        
        content_frame = ttk.Frame(dialog, style='Surface.TFrame', padding=20)
        content_frame.pack(fill='both', expand=True)
        
        ttk.Label(content_frame, text='Количество конфигураций:', style='Normal.TLabel').pack(anchor='w', pady=5)
        quantity_var = tk.IntVar(value=1)
        quantity_spin = ttk.Spinbox(content_frame, from_=1, to=100, textvariable=quantity_var, width=10, style='Modern.TEntry')
        quantity_spin.pack(fill='x', pady=5)
        
        def create_order():
            """Создание заказа"""
            quantity = quantity_var.get()
            
            try:
                headers = {'token': self.token}
                response = requests.post(
                    f'{self.base_url}/orders/create_order/',
                    headers=headers,
                    json={'configuration_id': self.current_config_id, 'quantity': quantity}
                )
                
                if response.status_code == 200:
                    result = response.json()
                    messagebox.showinfo('Успех!', f'Заказ создан! ID: {result["order_id"]}')

                    self.load_components()
                    
                    dialog.destroy()
                else:
                    error = response.json().get('detail', 'Ошибка создания заказа')
                    messagebox.showerror('Ошибка!', error)
            except requests.exceptions.RequestException as e:
                messagebox.showerror('Ошибка!', f'Ошибка соединения: {e}')
        
        button_frame = ttk.Frame(content_frame, style='Surface.TFrame')
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame, text='Создать заказ', command=create_order, 
                style='Primary.TButton').pack(side='left', padx=10)
        ttk.Button(button_frame, text='Отмена', command=dialog.destroy, style='Secondary.TButton').pack(side='left', padx=10)

    def delete_configuration(self):
        """Удаляет конфигурацию"""
        if not hasattr(self, 'current_config_id') or not self.current_config_id:
            messagebox.showwarning('Внимание!', 'Выберите конфигурацию для удаления')
            return
        
        config_name = self.configurations_data[self.current_config_id]['name_config']
        
        if messagebox.askyesno('Подтверждение', f'Удалить конфигурацию "{config_name}"?'):
            try:
                headers = {'token': self.token}
                response = requests.delete(
                    f'{self.base_url}/configurations/delete_by_id/',
                    headers=headers,
                    params={'config_id': self.current_config_id}
                )
                
                if response.status_code == 200:
                    messagebox.showinfo('Успех!', 'Конфигурация удалена')
                    self.current_config_id = None
                    self.load_configurations()
                    for item in self.config_components_tree.get_children():
                        self.config_components_tree.delete(item)
                    self.total_label.config(text='Общая сумма: 0 руб.')
                else:
                    error = response.json().get('detail', 'Ошибка удаления')
                    messagebox.showerror('Ошибка!', error)
            except requests.exceptions.RequestException as e:
                messagebox.showerror('Ошибка!', f'Ошибка соединения: {e}')
        
    def init_orders_tab(self):
        """Инициализирует вкладку заказов"""
        main_frame = ttk.Frame(self.tab_orders, style='Surface.TFrame')
        main_frame.pack(fill='both', expand=True)
        
        ttk.Label(main_frame, text='Мои заказы', style='Header.TLabel').pack(anchor='w', pady=(0, 20))

        button_frame = ttk.Frame(main_frame, style='Surface.TFrame')
        button_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Button(button_frame, text='Обновить', 
                command=self.load_orders, style='Secondary.TButton').pack(side='left')

        paned_window = ttk.PanedWindow(main_frame, orient=tk.HORIZONTAL)
        paned_window.pack(fill='both', expand=True, pady=10)

        left_frame = ttk.Frame(paned_window, style='Surface.TFrame', padding=10)
        paned_window.add(left_frame, weight=1)
        
        ttk.Label(left_frame, text='Список заказов:', style='Normal.TLabel').pack(anchor='w', pady=(0, 10))

        columns = ('id', 'date', 'total', 'status')
        self.orders_tree = ttk.Treeview(left_frame, columns=columns, show='headings', height=15)
        
        self.orders_tree.heading('id', text='ID заказа')
        self.orders_tree.heading('date', text='Дата заказа')
        self.orders_tree.heading('total', text='Общая сумма')
        self.orders_tree.heading('status', text='Статус')
        
        self.orders_tree.column('id', width=80)
        self.orders_tree.column('date', width=120)
        self.orders_tree.column('total', width=100)
        self.orders_tree.column('status', width=120)
        
        scrollbar_left = ttk.Scrollbar(left_frame, orient=tk.VERTICAL, command=self.orders_tree.yview)
        self.orders_tree.configure(yscroll=scrollbar_left.set)
        
        self.orders_tree.pack(side='left', fill='both', expand=True)
        scrollbar_left.pack(side='right', fill='y')

        right_frame = ttk.Frame(paned_window, style='Surface.TFrame', padding=10)
        paned_window.add(right_frame, weight=2)
        
        ttk.Label(right_frame, text='Детали заказа:', style='Normal.TLabel').pack(anchor='w', pady=(0, 10))

        order_columns = ('config_name', 'quantity', 'price', 'total')
        self.order_details_tree = ttk.Treeview(right_frame, columns=order_columns, show='headings', height=10)
        
        self.order_details_tree.heading('config_name', text='Конфигурация')
        self.order_details_tree.heading('quantity', text='Количество')
        self.order_details_tree.heading('price', text='Цена за шт.')
        self.order_details_tree.heading('total', text='Итого')
        
        self.order_details_tree.column('config_name', width=200)
        self.order_details_tree.column('quantity', width=80)
        self.order_details_tree.column('price', width=100)
        self.order_details_tree.column('total', width=100)
        
        scrollbar_right = ttk.Scrollbar(right_frame, orient=tk.VERTICAL, command=self.order_details_tree.yview)
        self.order_details_tree.configure(yscroll=scrollbar_right.set)
        
        self.order_details_tree.pack(fill='both', expand=True)
        scrollbar_right.pack(side='right', fill='y')

        info_frame = ttk.Frame(right_frame, style='Surface.TFrame')
        info_frame.pack(fill='x', pady=10)
        
        button_frame_right = ttk.Frame(right_frame, style='Surface.TFrame')
        button_frame_right.pack(fill='x', pady=10)

        ttk.Button(button_frame_right, text='Экспорт счета',
                   command=self.export_order_on_pdf, style='Primary.TButton').pack(side='right', padx=10)
        
        ttk.Button(button_frame_right, text='Оплатить',
                   command=self.pay_selected_order, style='Primary.TButton').pack(side='right', padx=10)
        
        ttk.Button(button_frame_right, text='Отменить заказ', 
                command=self.cancel_selected_order, style='Secondary.TButton').pack(side='right', padx=10)
        
        self.order_info_label = ttk.Label(info_frame, text='Выберите заказ для просмотра деталей', 
                                        style='Secondary.TLabel')
        self.order_info_label.pack(anchor='w', padx=10)

        self.orders_tree.bind('<<TreeviewSelect>>', self.on_order_select)

        self.load_orders()
    
    def amount_to_words(self, amount):
        """Преобразует сумму в рублях в слова"""
        try:
            rub = int(amount)
            kop = int((amount - rub) * 100)
            
            if rub == 0:
                rub_text = "ноль"
            else:
                rub_text = f"{rub}"
                
            result = f"{rub_text} руб. {kop:02d} коп."
            return result
            
        except:
            return f"{amount:.2f} руб."

    def export_order_on_pdf(self):
        """Экспорт выбранного заказа в формат PDF"""
        selected = self.orders_tree.selection()
        if not selected:
            messagebox.showwarning('Внимание!', 'Для осуществления экспорта выберите заказ')
            return
        
        try:
            try:
                pdfmetrics.registerFont(TTFont('Arial', 'C:/Windows/Fonts/arial.ttf'))
                pdfmetrics.registerFont(TTFont('Arial-Bold', 'C:/Windows/Fonts/arialbd.ttf'))
                font_normal = 'Arial'
                font_bold = 'Arial-Bold'
            except:
                try:
                    pdfmetrics.registerFont(TTFont('Arial', 'arial.ttf'))
                    pdfmetrics.registerFont(TTFont('Arial-Bold', 'arialbd.ttf'))
                    font_normal = 'Arial'
                    font_bold = 'Arial-Bold'
                except:
                    font_normal = 'Helvetica'
                    font_bold = 'Helvetica-Bold'
                    print("Предупреждение: Используются стандартные шрифты без поддержки кириллицы")

            item = selected[0]
            order_id = self.orders_tree.item(item)['values'][0]
            order_data = self.orders_data.get(order_id)

            if not order_data:
                messagebox.showerror('Ошибка!', 'Не удалось получить данные выбранного заказа')
                return
            
            os.makedirs('./Экспорт', exist_ok=True)
            safe_filename = f'Счет_ANTech_#{order_id}'
            file_path = f'./Экспорт/{safe_filename}.pdf'

            c = canvas.Canvas(file_path, pagesize=A4)
            width, height = A4

            c.setFont(font_bold, 18)
            c.drawString(50, height-50, "ANTech - Конфигуратор ПК")

            c.setFont(font_normal, 10)
            c.drawString(50, height-65, "Счет на оплату")

            c.setFont(font_bold, 12)
            c.drawString(400, height - 50, f"Счет №: {order_id}")
            order_date = order_data.get('order_date', '')
            if order_date:
                try:
                    formatted_date = order_date.split('T')[0]
                    c.setFont(font_normal, 10)
                    c.drawString(400, height - 65, f"Дата: {formatted_date}")
                except:
                    pass
            
            c.setFont(font_bold, 12)
            c.drawString(50, height - 100, "Покупатель:")
            c.setFont(font_normal, 10)
            
            y_position = height - 115
            user_info = [
                f"ФИО: {self.user_data.get('name', 'Не указано')}",
                f"Email: {self.user_data.get('email', 'Не указан')}",
                f"Телефон: {self.user_data.get('phone', 'Не указан')}",
                f"Адрес: {self.user_data.get('address', 'Не указан')}"
            ]
            
            for info in user_info:
                c.drawString(50, y_position, info)
                y_position -= 15
            
            y_position -= 20 
            
            c.setFont(font_bold, 10)
            c.drawString(50, y_position, "№")
            c.drawString(70, y_position, "Наименование")
            c.drawString(300, y_position, "Кол-во")
            c.drawString(350, y_position, "Цена за ед.")
            c.drawString(450, y_position, "Сумма")
            
            c.line(50, y_position - 5, 550, y_position - 5)
        
            y_position -= 20
            total_amount = 0
            configurations = order_data.get('configurations', [])
            
            c.setFont(font_normal, 9)
            for i, config in enumerate(configurations, 1):
                config_name = config.get('configuration_name', 'Конфигурация ПК')
                quantity = config.get('quantity', 1)
                price = config.get('price_at_time', 0)
                total = config.get('total', price * quantity)
                total_amount += total
                
                c.drawString(50, y_position, str(i))
                if len(config_name) > 40:
                    config_name = config_name[:37] + "..."
                c.drawString(70, y_position, config_name)
                c.drawString(300, y_position, str(quantity))
                c.drawString(350, y_position, f"{price:.2f} руб.")
                c.drawString(450, y_position, f"{total:.2f} руб.")
                
                y_position -= 15
                
                if y_position < 150:
                    c.showPage()
                    y_position = height - 50
                    c.setFont(font_normal, 9)

            c.line(50, y_position - 5, 550, y_position - 5)

            y_position -= 20
            c.setFont(font_bold, 12)
            c.drawString(350, y_position, f"ИТОГО: {total_amount:.2f} руб.")

            y_position -= 20
            c.setFont(font_normal, 9)
            amount_words = self.amount_to_words(total_amount)
            c.drawString(50, y_position, f"Всего к оплате: {amount_words}")

            y_position -= 20
            status = order_data.get('status_name', 'Неизвестно')
            c.setFont(font_bold, 10)
            c.drawString(50, y_position, f"Статус заказа: {status}")

            y_position -= 50

            c.setFont(font_normal, 10)
            c.drawString(50, y_position, "Подпись покупателя: ___________________")
            c.drawString(50, y_position - 15, "Расшифровка подписи: ___________________")

            c.drawString(350, y_position, "Подпись исполнителя: ___________________")
            c.drawString(350, y_position - 15, "М.П.")

            y_position -= 50
            c.setFont(font_bold, 10)
            c.drawString(50, y_position, "Контактная информация:")
            c.setFont(font_normal, 9)
            
            contact_info = [
                "ANTech - Конфигуратор ПК",
                "Email: support@antech.ru",
                "Телефон: +7 (777) 777-77-77", 
                "Адрес: г. Кострома",
                f"Документ сформирован: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}"
            ]
            
            for info in contact_info:
                y_position -= 12
                c.drawString(50, y_position, info)
        
            c.save()
            
            messagebox.showinfo('Успех!', f'Счет по заказу #{order_id} успешно экспортирован в PDF\nФайл: {safe_filename}.pdf')

        except Exception as e:
            messagebox.showerror('Ошибка!', f'Ошибка при экспорте в PDF: {str(e)}')

    def pay_selected_order(self):
        """Обрабатывает оплату выбранного заказа"""
        selected = self.orders_tree.selection()
        if not selected:
            messagebox.showwarning('Внимание!', 'Выберите заказ для оплаты')
            return
        
        item = selected[0]
        order_id = self.orders_tree.item(item)['values'][0]
        order_total = self.orders_tree.item(item)['values'][2]
        order_status = self.orders_tree.item(item)['values'][3]
        
        if order_status == 'Оплачен':
            messagebox.showerror('Ошибка!', 'Выбранная конфигурация уже оплачена')
            return

        pay_window = tk.Toplevel(self.root)
        pay_window.title('Оплата заказа')
        pay_window.geometry('1000x1000')
        pay_window.resizable(False, False)
        pay_window.configure(bg=self.style.background_color)
        pay_window.transient(self.root)
        pay_window.grab_set()

        pay_window.geometry(f'+{self.root.winfo_x() + 200}+{self.root.winfo_y() + 50}')

        title_font = ('Arial', 20, 'bold')
        label_font = ('Arial', 11)
        entry_font = ('Arial', 12)
        button_font = ('Arial', 12, 'bold')
        card_font = ('Arial', 16, 'bold')
        card_small_font = ('Arial', 10)

        content_frame = ttk.Frame(pay_window, style='Surface.TFrame', padding=30)
        content_frame.pack(fill='both', expand=True, padx=20, pady=20)

        ttk.Label(content_frame, text='Оплата заказа', font=title_font, 
                background=self.style.surface_color, foreground=self.style.text_primary).pack(pady=20)
 
        ttk.Label(content_frame, text=f'Заказ №{order_id}', font=('Arial', 14), 
                background=self.style.surface_color, foreground=self.style.success_color).pack(pady=5)
        ttk.Label(content_frame, text=f'Сумма: {order_total}', font=('Arial', 12), 
                background=self.style.surface_color, foreground=self.style.text_primary).pack(pady=5)
        
        card_frame = tk.Frame(content_frame, bg=self.style.primary_color, bd=2, relief='raised', 
                            width=400, height=220)
        card_frame.pack(pady=20, padx=50, fill='both')
        card_frame.pack_propagate(False)

        tk.Label(card_frame, text='МИР', font=('Arial', 24, 'bold'), 
                bg=self.style.primary_color, fg='white').place(x=610, y=20)

        tk.Label(card_frame, text='Номер карты', font=card_small_font, 
                bg=self.style.primary_color, fg='#ecf0f1', anchor='w').place(x=30, y=80)
        card_display = tk.Label(card_frame, text='**** **** **** ****', 
                            font=card_font, bg=self.style.primary_color, fg='white')
        card_display.place(x=30, y=100)

        tk.Label(card_frame, text='Срок действия', font=card_small_font, 
                bg=self.style.primary_color, fg='#ecf0f1').place(x=30, y=150)
        tk.Label(card_frame, text='MM / ГГ', font=('Arial', 12), 
                bg=self.style.primary_color, fg='white').place(x=30, y=170)
        
        tk.Label(card_frame, text='CVC / CVV', font=card_small_font, 
                bg=self.style.primary_color, fg='#ecf0f1').place(x=150, y=150)
        tk.Label(card_frame, text='***', font=('Arial', 12), 
                bg=self.style.primary_color, fg='white').place(x=150, y=170)

        form_frame = tk.Frame(content_frame, bg=self.style.surface_color)
        form_frame.pack(pady=20, padx=50, fill='both')

        ttk.Label(form_frame, text='Номер карты:', font=label_font, 
                background=self.style.surface_color, foreground=self.style.text_primary).pack(anchor='w', pady=5)
        
        card_number_var = tk.StringVar()
        card_number_entry = ttk.Entry(form_frame, textvariable=card_number_var, 
                                    font=entry_font, width=25, style='Modern.TEntry')
        card_number_entry.pack(fill='x', pady=5)

        expiry_cvc_frame = tk.Frame(form_frame, bg=self.style.surface_color)
        expiry_cvc_frame.pack(fill='x', pady=10)

        expiry_frame = tk.Frame(expiry_cvc_frame, bg=self.style.surface_color)
        expiry_frame.pack(side='left', padx=(0, 20))
        
        ttk.Label(expiry_frame, text='Срок действия (ММ/ГГ):', font=label_font, 
                background=self.style.surface_color, foreground=self.style.text_primary).pack(anchor='w')
        
        expiry_subframe = tk.Frame(expiry_frame, bg=self.style.surface_color)
        expiry_subframe.pack(fill='x', pady=5)
        
        month_var = tk.StringVar()
        month_entry = ttk.Entry(expiry_subframe, textvariable=month_var, 
                            font=entry_font, width=5, style='Modern.TEntry')
        month_entry.pack(side='left', padx=(0, 5))
        
        ttk.Label(expiry_subframe, text='/', font=entry_font, 
                background=self.style.surface_color, foreground=self.style.text_primary).pack(side='left')
        
        year_var = tk.StringVar()
        year_entry = ttk.Entry(expiry_subframe, textvariable=year_var, 
                            font=entry_font, width=5, style='Modern.TEntry')
        year_entry.pack(side='left', padx=(5, 0))
 
        cvc_frame = tk.Frame(expiry_cvc_frame, bg=self.style.surface_color)
        cvc_frame.pack(side='left')
        
        ttk.Label(cvc_frame, text='CVC/CVV код:', font=label_font, 
                background=self.style.surface_color, foreground=self.style.text_primary).pack(anchor='w')
        
        cvc_var = tk.StringVar()
        cvc_entry = ttk.Entry(cvc_frame, textvariable=cvc_var, 
                            font=entry_font, width=8, show='*', style='Modern.TEntry')
        cvc_entry.pack(fill='x', pady=5)

        ttk.Label(form_frame, text='Имя владельца карты:', font=label_font, 
                background=self.style.surface_color, foreground=self.style.text_primary).pack(anchor='w', pady=5)
        
        name_var = tk.StringVar()
        name_entry = ttk.Entry(form_frame, textvariable=name_var, 
                            font=entry_font, width=25, style='Modern.TEntry')
        name_entry.pack(fill='x', pady=5)

        def update_card_display(*args):
            card_number = card_number_var.get().replace(' ', '')

            if len(card_number) == 0:
                displayed = '*' * 16
            else:
                displayed = card_number.ljust(16, '*')[:16]

            formatted = ' '.join([displayed[i:i+4] for i in range(0, 16, 4)])
            card_display.config(text=formatted)
        
        def update_expiry_display(*args):
            month = month_var.get().ljust(2, '*')
            year = year_var.get().ljust(2, '*')
        
        def update_cvc_display(*args):
            cvc = cvc_var.get().ljust(3, '*')

        card_number_var.trace('w', update_card_display)
        month_var.trace('w', update_expiry_display)
        year_var.trace('w', update_expiry_display)
        cvc_var.trace('w', update_cvc_display)

        def format_card_number(*args):
            value = card_number_var.get().replace(' ', '')[:16]
            formatted = ' '.join([value[i:i+4] for i in range(0, len(value), 4)])
            if card_number_var.get() != formatted:
                card_number_var.set(formatted)
        
        card_number_var.trace('w', format_card_number)

        def validate_month(*args):
            value = month_var.get()
            if value.isdigit() and len(value) <= 2:
                if value and int(value) > 12:
                    month_var.set('12')
                return True
            elif value == '':
                return True
            else:
                month_var.set(''.join(filter(str.isdigit, value)))
        
        def validate_year(*args):
            value = year_var.get()
            if value.isdigit() and len(value) <= 2:
                return True
            elif value == '':
                return True
            else:
                year_var.set(''.join(filter(str.isdigit, value)))
        
        def validate_cvc(*args):
            value = cvc_var.get()
            if value.isdigit() and len(value) <= 3:
                return True
            elif value == '':
                return True
            else:
                cvc_var.set(''.join(filter(str.isdigit, value)))
        
        month_var.trace('w', lambda *args: validate_month())
        year_var.trace('w', lambda *args: validate_year())
        cvc_var.trace('w', lambda *args: validate_cvc())

        button_frame = tk.Frame(content_frame, bg=self.style.surface_color)
        button_frame.pack(pady=20)
        
        def process_payment():
            """Обработка оплаты"""
            if not all([card_number_var.get(), month_var.get(), 
                    year_var.get(), cvc_var.get(), name_var.get()]):
                messagebox.showerror('Ошибка!', 'Заполните все поля')
                return

            card_number = card_number_var.get().replace(' ', '')
            if len(card_number) != 16 or not card_number.isdigit():
                messagebox.showerror('Ошибка!', 'Введите корректный номер карты (16 цифр)')
                return

            if len(month_var.get()) != 2 or not month_var.get().isdigit():
                messagebox.showerror('Ошибка!', 'Введите корректный месяц (2 цифры)')
                return
            
            month = int(month_var.get())
            if month < 1 or month > 12:
                messagebox.showerror('Ошибка!', 'Месяц должен быть от 01 до 12')
                return
            
            if len(year_var.get()) != 2 or not year_var.get().isdigit():
                messagebox.showerror('Ошибка!', 'Введите корректный год (2 цифры)')
                return

            if len(cvc_var.get()) != 3 or not cvc_var.get().isdigit():
                messagebox.showerror('Ошибка!', 'Введите корректный CVC код (3 цифры)')
                return

            try:
                headers = {'token': self.token}
 
                update_response = requests.put(
                    f'{self.base_url}/orders/user/update_order_status/',
                    headers=headers,
                    params={
                        'order_id': order_id,
                        'new_status': 'Оплачен'
                    }
                )
                
                if update_response.status_code == 200:
                    messagebox.showinfo('Успех!', 
                        f'Заказ #{order_id} успешно оплачен!\n'
                        f'Сумма: {order_total}\n'
                        f'Спасибо за покупку!')
                    pay_window.destroy()
                    self.load_orders()
                else:
                    error = update_response.json().get('detail', 'Ошибка обновления статуса')
                    messagebox.showerror('Ошибка!', error)
                            
            except requests.exceptions.RequestException as e:
                messagebox.showerror('Ошибка!', f'Ошибка соединения: {e}')
        
        ttk.Button(button_frame, text='Оплатить', 
                command=process_payment, style='Primary.TButton').pack(side='left', padx=10)
        
        ttk.Button(button_frame, text='Отмена', 
                command=pay_window.destroy, style='Secondary.TButton').pack(side='left', padx=10)

        card_number_entry.focus()

        expiry_display = tk.Label(card_frame, text='MM / ГГ', font=('Arial', 12), 
                                bg=self.style.primary_color, fg='white')
        expiry_display.place(x=30, y=170)
        
        cvc_display = tk.Label(card_frame, text='***', font=('Arial', 12), 
                            bg=self.style.primary_color, fg='white')
        cvc_display.place(x=150, y=170)
        
        name_display = tk.Label(card_frame, text='ВЛАДЕЛЕЦ КАРТЫ', font=('Arial', 10), 
                            bg=self.style.primary_color, fg='white')
        name_display.place(x=30, y=190)

        def update_expiry_display(*args):
            month = month_var.get().ljust(2, '*')[:2]
            year = year_var.get().ljust(2, '*')[:2]
            expiry_display.config(text=f'{month} / {year}')
        
        def update_cvc_display(*args):
            cvc = cvc_var.get().ljust(3, '*')[:3]
            cvc_display.config(text=cvc)
        
        def update_name_display(*args):
            name = name_var.get().upper() or 'ВЛАДЕЛЕЦ КАРТЫ'
            name_display.config(text=name[:20])
        
        month_var.trace('w', update_expiry_display)
        year_var.trace('w', update_expiry_display)
        cvc_var.trace('w', update_cvc_display)
        name_var.trace('w', update_name_display)
            
    def cancel_selected_order(self):
        """Отменяет выбранный заказ"""
        selected = self.orders_tree.selection()
        if not selected:
            messagebox.showwarning('Внимание!', 'Выберите заказ для отмены')
            return
        
        item = selected[0]
        order_id = self.orders_tree.item(item)['values'][0]
        
        if messagebox.askyesno('Подтверждение', f'Вы уверены, что хотите отменить заказ #{order_id}?'):
            try:
                headers = {'token': self.token}
                response = requests.delete(
                    f'{self.base_url}/orders/cancel_order/',
                    headers=headers,
                    params={'order_id': order_id}
                )
                
                if response.status_code == 200:
                    messagebox.showinfo('Успех!', 'Заказ отменен')
                    self.load_orders()
                    for item in self.order_details_tree.get_children():
                        self.order_details_tree.delete(item)
                    self.order_info_label.config(text='Выберите заказ для просмотра деталей')
                else:
                    error = response.json().get('detail', 'Ошибка отмены заказа')
                    messagebox.showerror('Ошибка!', error)
            except requests.exceptions.RequestException as e:
                messagebox.showerror('Ошибка!', f'Ошибка соединения: {e}')

    def load_orders(self):
        """Загружает список заказов"""
        try:
            headers = {'token': self.token}
            response = requests.get(f'{self.base_url}/orders/get_user_orders/', headers=headers)
            
            if response.status_code == 200:
                for item in self.orders_tree.get_children():
                    self.orders_tree.delete(item)
                
                orders = response.json()
                self.orders_data = {}
                
                for order in orders:
                    order_id = order['id']
                    self.orders_data[order_id] = order

                    order_date = order.get('order_date', '')
                    if order_date:
                        try:
                            order_date = order_date.split('T')[0]
                        except:
                            pass

                    total_amount = order.get('total_amount', 0)
                    total_text = f"{total_amount:.2f} руб." if total_amount else "0 руб."
                    
                    self.orders_tree.insert('', 'end', values=(
                        order_id,
                        order_date,
                        total_text,
                        order.get('status_name', 'Неизвестно')
                    ))
            else:
                messagebox.showerror('Ошибка!', 'Не удалось загрузить заказы')
        except requests.exceptions.RequestException as e:
            messagebox.showerror('Ошибка!', f'Ошибка соединения: {e}')

    def on_order_select(self, event):
        """Обрабатывает выбор заказа"""
        selected = self.orders_tree.selection()
        if not selected:
            return
        
        item = selected[0]
        order_id = self.orders_tree.item(item)['values'][0]
        
        if order_id in self.orders_data:
            order = self.orders_data[order_id]
            self.display_order_details(order)

    def display_order_details(self, order):
        """Отображает детали заказа"""
        for item in self.order_details_tree.get_children():
            self.order_details_tree.delete(item)

        order_id = order['id']
        order_date = order.get('order_date', 'Неизвестно')
        total_amount = order.get('total_amount', 0)
        status = order.get('status_name', 'Неизвестно')
        
        info_text = f"Заказ #{order_id} | Дата: {order_date} | Статус: {status} | Общая сумма: {total_amount:.2f} руб."
        self.order_info_label.config(text=info_text)

        configurations = order.get('configurations', [])
        for config in configurations:
            config_name = config.get('configuration_name', 'Неизвестно')
            quantity = config.get('quantity', 1)
            price = config.get('price_at_time', 0)
            total = config.get('total', price * quantity)
            
            self.order_details_tree.insert('', 'end', values=(
                config_name,
                quantity,
                f"{price:.2f} руб.",
                f"{total:.2f} руб."
            ))
        
    def init_profile_tab(self):
        """Инициализирует вкладку профиля"""
        profile_frame = ttk.Frame(self.tab_profile, style='Surface.TFrame')
        profile_frame.pack(fill='both', expand=True)
        
        ttk.Label(profile_frame, text='Мой профиль', style='Header.TLabel').pack(anchor='w', pady=(0, 20))
        
        info_frame = ttk.LabelFrame(profile_frame, text='Информация о профиле', style='Surface.TFrame', padding=20)
        info_frame.pack(fill='x', pady=10, padx=10)
        
        user_info = [
            ('ФИО:', self.user_data['name']),
            ('Email:', self.user_data['email']),
            ('Телефон:', self.user_data['phone']),
            ('Роль:', self.user_data['role']),
            ('Адрес:', self.user_data['address'] or 'Не указан')
        ]
        
        for i, (label, value) in enumerate(user_info):
            row_frame = ttk.Frame(info_frame, style='Surface.TFrame')
            row_frame.pack(fill='x', pady=5)
            
            ttk.Label(row_frame, text=label, style='Normal.TLabel', width=15, anchor='e').pack(side='left', padx=5)
            ttk.Label(row_frame, text=value, style='Normal.TLabel', anchor='w').pack(side='left', padx=5, fill='x', expand=True)
            
        button_frame = ttk.Frame(profile_frame, style='Surface.TFrame')
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame, text='Изменить адрес',
                   command=self.edit_address, style='Primary.TButton').pack(side='left', padx=10)
        
        ttk.Button(button_frame, text='Обновить данные',
                   command=self.load_user_data, style='Secondary.TButton').pack(side='left', padx=10)
        
        ttk.Button(button_frame, text='Сменить пароль', command=self.change_password,
                   style='Secondary.TButton').pack(side='left', padx=10)
        
        ttk.Button(button_frame, text='Выйти из аккаунта',
                   command=self.logout, style='Secondary.TButton').pack(side='left', padx=10)
        
    def edit_address(self):
        """Редактирует адрес пользователя"""
        edit_window = tk.Toplevel(self.root)
        edit_window.title('Изменение адреса')
        edit_window.geometry('600x500')
        edit_window.resizable(False, False)
        edit_window.configure(bg=self.style.background_color)
        edit_window.grab_set()
        edit_window.transient(self.root)
        edit_window.geometry(f'+{self.root.winfo_x() + 100}+{self.root.winfo_y() + 100}')
        
        content_frame = ttk.Frame(edit_window, style='Surface.TFrame', padding=20)
        content_frame.pack(fill='both', expand=True)
        
        ttk.Label(content_frame, text='Изменение адреса доставки', style='Header.TLabel').pack(pady=20)
        
        current_address = self.user_data.get('address', 'Не указан')
        ttk.Label(content_frame, text='Новый адрес:', style='Normal.TLabel').pack(anchor='w', pady=5)
        address_entry = ttk.Entry(content_frame, width=50, font=('Arial', 12), style='Modern.TEntry')
        address_entry.pack(fill='x', pady=5)
        
        def save_address():
            new_address = address_entry.get().strip()
            if not new_address:
                messagebox.showerror('Ошибка!', 'Поле адреса не может быть пустым')
                return
            
            try:
                token = self.token
                if not token:
                    messagebox.showerror('Ошибка!', 'Токен авторизации не найден')
                    return
                
                response = requests.put(
                    f'{self.base_url}/users/edit_address/',
                    params={'new_address': new_address},
                    headers={'token': self.token}
                )
                
                if response.status_code == 200:
                    messagebox.showinfo('Успех!', 'Адрес успешно изменен')
                    self.user_data['address'] = new_address
                    edit_window.destroy()
                else:
                    error = response.json().get('message', 'Не удалось изменить адрес')
                    messagebox.showerror('Ошибка!', error)
            except requests.exceptions.RequestException as e:
                messagebox.showerror('Ошибка!', f'Не удалось подключиться к серверу: {e}')
        
        button_frame = ttk.Frame(content_frame, style='Surface.TFrame')
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame, text='Сохранить', command=save_address, style='Primary.TButton').pack(side='left', padx=10)
        ttk.Button(button_frame, text='Отмена', command=edit_window.destroy, style='Secondary.TButton').pack(side='left', padx=10)
    
    def change_password(self):
        """Изменяет пароль пользователя"""
        password_window = tk.Toplevel(self.root)
        password_window.title('Смена пароля')
        password_window.geometry('600x500')
        password_window.resizable(False, False)
        password_window.configure(bg=self.style.background_color)
        password_window.grab_set()
        
        password_window.transient(self.root)
        password_window.geometry(f'+{self.root.winfo_x() + 100}+{self.root.winfo_y() + 100}')
        
        user_email = self.user_data['email']
        content_frame = ttk.Frame(password_window, style='Surface.TFrame', padding=20)
        content_frame.pack(fill='both', expand=True)
        
        ttk.Label(content_frame, text='Смена пароля', style='Header.TLabel').pack(pady=20)
        ttk.Label(content_frame, text=f'Для смены пароля будет отправлен код на email: {user_email}', style='Normal.TLabel', wraplength=400).pack(pady=10)
        
        code_frame = ttk.Frame(content_frame, style='Surface.TFrame')

        def show_code_entry():
            send_btn.pack_forget()
            
            ttk.Label(code_frame, text='Код подтверждения', style='Normal.TLabel').pack(anchor='w', pady=5)
            code_entry = ttk.Entry(code_frame, width=30, font=('Arial', 12), style='Modern.TEntry')
            code_entry.pack(fill='x', pady=5)
            
            ttk.Label(code_frame, text='Новый пароль:', style='Normal.TLabel').pack(anchor='w', pady=5)
            new_password_entry = ttk.Entry(code_frame, width=30, show='*', font=('Arial', 12), style='Modern.TEntry')
            new_password_entry.pack(fill='x', pady=5)
            
            ttk.Label(code_frame, text='Подтвердите пароль:', style='Normal.TLabel').pack(anchor='w', pady=5)
            confirm_password_entry = ttk.Entry(code_frame, width=30, show='*', font=('Arial', 12), style='Modern.TEntry')
            confirm_password_entry.pack(fill='x', pady=5)
            
            def confirm_change():
                code = code_entry.get().strip()
                new_password = new_password_entry.get().strip()
                confirm_password = confirm_password_entry.get().strip()
                
                if not all([code, new_password, confirm_password]):
                    messagebox.showerror('Ошибка!', 'Заполните все поля')
                    return
                
                if new_password != confirm_password:
                    messagebox.showerror('Ошибка!', 'Пароли не совпадают')
                    return
                
                try:
                    response = requests.post(
                        f'{self.base_url}/users/confirm_change_password/',
                        params={
                            'email': user_email,
                            'code': code,
                            'new_password': new_password
                        }
                    )
                    
                    if response.status_code == 200:
                        messagebox.showinfo('Успех!', 'Пароль успешно изменен')
                        password_window.destroy()
                    else:
                        error = response.json().get('detail', 'Ошибка смены пароля')
                        messagebox.showerror('Ошибка', error)
                except requests.exceptions.RequestException as e:
                    messagebox.showerror('Ошибка!', f'Не удалось подключиться к серверу: {e}')
            
            ttk.Button(code_frame, text='Сменить пароль', command=confirm_change, style='Primary.TButton').pack(pady=20)
            code_frame.pack()
        
        def send_code():
            try:
                response = requests.post(
                    f'{self.base_url}/users/change_password/',
                    params={'email': user_email}
                )
                
                if response.status_code == 200:
                    messagebox.showinfo('Успех!', 'Код подтверждения отправлен на ваш email')
                    show_code_entry()
                else:
                    error = response.json().get('detail', 'Не удалось отправить код подтверждения')
                    messagebox.showerror('Ошибка!', error)
            except requests.exceptions.RequestException as e:
                messagebox.showerror('Ошибка!', f'Не удалось подключиться к серверу: {e}')

        send_btn = ttk.Button(content_frame, text='Отправить код', command=send_code, style='Primary.TButton')
        send_btn.pack(pady=20)
        
        ttk.Button(content_frame, text='Отмена', command=password_window.destroy, style='Secondary.TButton').pack(pady=10)

class AdminApp:
    """Административная панель управления"""
    def __init__(self, root, token):
        self.root = root
        self.root.title('ANTech - Панель администратора')
        self.root.state('zoomed')
        self.style = ModernStyle()
        self.setup_styles()
        self.root.configure(bg=self.style.background_color)
        self.token = token
        self.base_url = 'http://127.0.0.1:8000'
        
        self.main_container = ttk.Frame(root, style='Modern.TFrame')
        self.main_container.pack(fill='both', expand=True)
        
        self.create_header()
        self.create_tabs()
        
        self.init_users_tab()
        self.init_manufactures_tab()
        self.init_components_tab()
        self.init_configurations_tab()
        self.init_orders_tab()
    
    def setup_styles(self):
        """Настройка стилей виджетов"""
        style = ttk.Style()
        style.configure('Modern.TFrame', background=self.style.background_color)
        style.configure('Surface.TFrame', background=self.style.surface_color)
        style.configure('Header.TLabel', 
                       font=('Arial', 16, 'bold'),
                       background=self.style.surface_color,
                       foreground=self.style.text_primary)
        style.configure('Normal.TLabel', 
                       font=('Arial', 11),
                       background=self.style.surface_color,
                       foreground=self.style.text_primary)
        style.configure('Secondary.TLabel',
                       font=('Arial', 11),
                       background=self.style.surface_color,
                       foreground=self.style.text_secondary)
        style.configure('Primary.TButton', 
                       font=('Arial', 11), 
                       background=self.style.primary_color,
                       foreground='#1e3d6d',
                       borderwidth=0,
                       focuscolor='none')
        style.map('Primary.TButton',
                 background=[('active', self.style.secondary_color)])
        style.configure('Secondary.TButton',
                       font=('Arial', 11),
                       background=self.style.surface_color,
                       foreground=self.style.primary_color,
                       borderwidth=1)
        style.map('Secondary.TButton',
                 background=[('active', self.style.background_color)])
        style.configure('Modern.TNotebook', 
                       background=self.style.background_color)
        style.configure('Modern.TNotebook.Tab', 
                       background=self.style.surface_color,
                       foreground=self.style.text_primary,
                       padding=(20, 10))
        style.map('Modern.TNotebook.Tab',
                 background=[('selected', self.style.primary_color),
                           ('active', self.style.secondary_color)],
                 foreground=[('selected', 'white')])
        style.configure('Treeview',
                       background=self.style.surface_color,
                       foreground=self.style.text_primary,
                       fieldbackground=self.style.surface_color,
                       rowheight=25)
        style.configure('Treeview.Heading',
                       background=self.style.primary_color,
                       foreground="#1e3d6d",
                       font=('Arial', 10, 'bold'))
        style.map('Treeview.Heading',
                 background=[('active', self.style.secondary_color)])
        style.configure('Modern.TEntry',
                       fieldbackground=self.style.surface_color,
                       foreground=self.style.text_primary,
                       borderwidth=1,
                       relief='solid')
        style.configure('Modern.TCombobox',
                       fieldbackground=self.style.surface_color,
                       foreground=self.style.text_primary)
    
    def create_header(self):
        """Создает заголовок административной панели"""
        header_frame = ttk.Frame(self.main_container, style='Surface.TFrame', padding=20)
        header_frame.pack(fill='x', padx=20, pady=10)
        
        title_frame = ttk.Frame(header_frame, style='Surface.TFrame')
        title_frame.pack(side='left')
        
        ttk.Label(title_frame, text='ANTech - Панель администратора', 
                 style='Header.TLabel').pack(anchor='w')
        ttk.Label(title_frame, text='Управление системой', 
                 style='Secondary.TLabel').pack(anchor='w')
        
        ttk.Button(header_frame, text='Выход', 
                  command=self.logout, style='Secondary.TButton').pack(side='right')
    
    def create_tabs(self):
        """Создает вкладки административной панели"""
        tab_control = ttk.Notebook(self.main_container, style='Modern.TNotebook')
        
        self.tab_users = ttk.Frame(tab_control, style='Surface.TFrame', padding=20)
        self.tab_manufactures = ttk.Frame(tab_control, style='Surface.TFrame', padding=20)
        self.tab_components = ttk.Frame(tab_control, style='Surface.TFrame', padding=20)
        self.tab_configurations = ttk.Frame(tab_control, style='Surface.TFrame', padding=20)
        self.tab_orders = ttk.Frame(tab_control, style='Surface.TFrame', padding=20)
        
        tab_control.add(self.tab_users, text='Пользователи')
        tab_control.add(self.tab_manufactures, text='Производители')
        tab_control.add(self.tab_components, text='Компоненты')
        tab_control.add(self.tab_configurations, text='Конфигурации')
        tab_control.add(self.tab_orders, text='Заказы')
        
        tab_control.pack(fill='both', expand=True, padx=20, pady=10)
    
    def logout(self):
        """Выполняет выход из системы"""
        for widget in self.root.winfo_children():
            widget.destroy()
        ModernAuthApp(self.root)
    
    def make_api_request(self, endpoint, method='GET', params=None, json_data=None):
        """Выполняет API запрос к серверу"""
        try:
            headers = {'token': self.token}
            url = f'{self.base_url}{endpoint}'
            
            if method == 'GET':
                response = requests.get(url, headers=headers, params=params)
            elif method == 'POST':
                response = requests.post(url, headers=headers, json=json_data, params=params)
            elif method == 'PUT':
                response = requests.put(url, headers=headers, json=json_data, params=params)
            elif method == 'DELETE':
                response = requests.delete(url, headers=headers, params=params)
            
            if response.status_code == 200:
                return response.json()
            else:
                error = response.json().get('detail', 'Ошибка API')
                messagebox.showerror('Ошибка!', error)
                return None
        except requests.exceptions.RequestException as e:
            messagebox.showerror('Ошибка!', f'Ошибка соединения: {e}')
            return None
    
    def init_users_tab(self):
        """Инициализирует вкладку пользователей"""
        main_frame = ttk.Frame(self.tab_users, style='Surface.TFrame')
        main_frame.pack(fill='both', expand=True)
        
        ttk.Label(main_frame, text='Управление пользователями', style='Header.TLabel').pack(anchor='w', pady=(0, 20))

        search_frame = ttk.Frame(main_frame, style='Surface.TFrame')
        search_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Label(search_frame, text='Поиск пользователя:', style='Normal.TLabel').pack(side='left', padx=5)
        
        self.user_search_var = tk.StringVar()
        search_entry = ttk.Entry(search_frame, textvariable=self.user_search_var, width=30, style='Modern.TEntry')
        search_entry.pack(side='left', padx=5)
        
        ttk.Button(search_frame, text='Найти', 
                  command=self.search_user, style='Primary.TButton').pack(side='left', padx=5)
        ttk.Button(search_frame, text='Обновить', 
                  command=self.load_users, style='Secondary.TButton').pack(side='left', padx=5)

        button_frame = ttk.Frame(main_frame, style='Surface.TFrame')
        button_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Button(button_frame, text='Изменить роль', 
                  command=self.change_user_role, style='Primary.TButton').pack(side='left', padx=5)

        columns = ('id', 'name', 'email', 'phone', 'role', 'address')
        self.users_tree = ttk.Treeview(main_frame, columns=columns, show='headings', height=20)
        
        self.users_tree.heading('id', text='ID')
        self.users_tree.heading('name', text='ФИО')
        self.users_tree.heading('email', text='Email')
        self.users_tree.heading('phone', text='Телефон')
        self.users_tree.heading('role', text='Роль')
        self.users_tree.heading('address', text='Адрес')
        
        self.users_tree.column('id', width=50)
        self.users_tree.column('name', width=150)
        self.users_tree.column('email', width=150)
        self.users_tree.column('phone', width=120)
        self.users_tree.column('role', width=100)
        self.users_tree.column('address', width=200)
        
        scrollbar = ttk.Scrollbar(main_frame, orient=tk.VERTICAL, command=self.users_tree.yview)
        self.users_tree.configure(yscroll=scrollbar.set)
        
        self.users_tree.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')

        search_entry.bind('<Return>', lambda event: self.search_user())
        
        self.load_users()
    
    def search_user(self):
        """Выполняет поиск пользователя"""
        search_text = self.user_search_var.get().strip()
        
        if not search_text:
            messagebox.showwarning('Внимание!', 'Введите email или номер телефона для поиска')
            return

        if self.is_phone(search_text):
            search_data = {'phone': search_text}
        else:
            search_data = {'email': search_text}
        
        try:
            headers = {'token': self.token}
            response = requests.post(
                f'{self.base_url}/users/get_user_by_email_or_phone/',
                headers=headers,
                json=search_data
            )
            
            if response.status_code == 200:
                user = response.json()

                for item in self.users_tree.get_children():
                    self.users_tree.delete(item)

                self.users_tree.insert('', 'end', values=(
                    user['id'],
                    user['name'],
                    user['email'],
                    user['phone'],
                    user['role'],
                    user['address']
                ))
                
                messagebox.showinfo('Успех!', f'Пользователь найден: {user["name"]}')
                
            else:
                error = response.json().get('detail', 'Пользователь не найден')
                messagebox.showerror('Ошибка!', error)
                
        except requests.exceptions.RequestException as e:
            messagebox.showerror('Ошибка!', f'Ошибка соединения: {e}')
        
    def is_phone(self, text):
        """Проверяет, является ли текст телефонным номером"""
        phone_pattern = r'^[\d\s\-\+\(\)]+$'
        return bool(re.match(phone_pattern, text)) and len(text) >= 5
    
    def load_users(self):
        """Загружает список пользователей"""
        data = self.make_api_request('/users/get_all/')
        if data:
            for item in self.users_tree.get_children():
                self.users_tree.delete(item)
            
            for user in data:
                self.users_tree.insert('', 'end', values=(
                    user['id'],
                    user['name'],
                    user['email'],
                    user['phone'],
                    user['role'],
                    user['address']
                ))
    
    def change_user_role(self):
        """Изменяет роль пользователя"""
        selected = self.users_tree.selection()
        if not selected:
            messagebox.showwarning('Внимание!', 'Выберите пользователя')
            return
        
        item = selected[0]
        user_data = self.users_tree.item(item)['values']
        user_email = user_data[2]
        current_role = user_data[4]
        
        dialog = tk.Toplevel(self.root)
        dialog.title('Изменение роли пользователя')
        dialog.geometry('350x200')
        dialog.configure(bg=self.style.background_color)
        dialog.transient(self.root)
        dialog.grab_set()
        
        content_frame = ttk.Frame(dialog, style='Surface.TFrame', padding=20)
        content_frame.pack(fill='both', expand=True)
        
        ttk.Label(content_frame, text=f'Пользователь: {user_email}', style='Normal.TLabel').pack(pady=10)
        ttk.Label(content_frame, text=f'Текущая роль: {current_role}', style='Normal.TLabel').pack(pady=5)
        
        role_var = tk.StringVar(value=current_role)
        role_combo = ttk.Combobox(content_frame, textvariable=role_var, state='readonly', style='Modern.TCombobox')
        role_combo['values'] = ('Пользователь', 'Администратор')
        role_combo.pack(pady=10)
        
        def save_role():
            new_role = role_var.get()
            if new_role == current_role:
                messagebox.showinfo('Информация', 'Роль не изменена')
                dialog.destroy()
                return
            
            data = self.make_api_request('/users/set_role/', method='POST',
                                       json_data={'email': user_email, 'new_role': new_role})
            if data:
                messagebox.showinfo('Успех!', 'Роль пользователя изменена')
                dialog.destroy()
                self.load_users()
        
        ttk.Button(content_frame, text='Сохранить', command=save_role, style='Primary.TButton').pack(pady=10)
        
    def init_manufactures_tab(self):
        """Инициализирует вкладку производителей"""
        main_frame = ttk.Frame(self.tab_manufactures, style='Surface.TFrame')
        main_frame.pack(fill='both', expand=True)
        
        ttk.Label(main_frame, text='Управление производителями', style='Header.TLabel').pack(anchor='w', pady=(0, 20))

        button_frame = ttk.Frame(main_frame, style='Surface.TFrame')
        button_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Button(button_frame, text='Добавить производителя', 
                  command=self.add_manufacture, style='Primary.TButton').pack(side='left', padx=5)
        ttk.Button(button_frame, text='Обновить', 
                  command=self.load_manufactures, style='Secondary.TButton').pack(side='left', padx=5)
        ttk.Button(button_frame, text='Удалить', 
                  command=self.delete_manufacture, style='Secondary.TButton').pack(side='left', padx=5)

        columns = ('id', 'name')
        self.manufactures_tree = ttk.Treeview(main_frame, columns=columns, show='headings', height=20)
        
        self.manufactures_tree.heading('id', text='ID')
        self.manufactures_tree.heading('name', text='Название')
        
        self.manufactures_tree.column('id', width=50)
        self.manufactures_tree.column('name', width=200)
        
        scrollbar = ttk.Scrollbar(main_frame, orient=tk.VERTICAL, command=self.manufactures_tree.yview)
        self.manufactures_tree.configure(yscroll=scrollbar.set)
        
        self.manufactures_tree.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')
        
        self.load_manufactures()
    
    def load_manufactures(self):
        """Загружает список производителей"""
        try:
            data = self.make_api_request('/manufactures/get_manufactures/')
            if data:
                for item in self.manufactures_tree.get_children():
                    self.manufactures_tree.delete(item)

                for manufacture in data:
                    self.manufactures_tree.insert('', 'end', values=(
                        manufacture.get('id', 'N/A'),
                        manufacture.get('name', 'Неизвестно')
                    ))
        except requests.exceptions.RequestException as e:
            messagebox.showerror('Ошибка!', f'Произошла ошибка при загрузке списка производителей: {e}')
      
    def add_manufacture(self):
        """Добавляет производителя"""
        dialog = tk.Toplevel(self.root)
        dialog.title('Добавление производителя')
        dialog.geometry('300x200')
        dialog.configure(bg=self.style.background_color)
        dialog.transient(self.root)
        dialog.grab_set()
        
        content_frame = ttk.Frame(dialog, style='Surface.TFrame', padding=20)
        content_frame.pack(fill='both', expand=True)
        
        ttk.Label(content_frame, text='Название производителя:', style='Normal.TLabel').pack(anchor='w', pady=5)
        name_entry = ttk.Entry(content_frame, width=30, style='Modern.TEntry')
        name_entry.pack(fill='x', pady=5)
        
        def save_manufacture():
            name = name_entry.get().strip()
            if not name:
                messagebox.showerror('Ошибка!', 'Введите название')
                return
            
            data = self.make_api_request('/manufactures/add_manufacture/', method='POST',
                                       params={'name': name})
            if data:
                messagebox.showinfo('Успех!', 'Производитель добавлен')
                dialog.destroy()
                self.load_manufactures()
        
        button_frame = ttk.Frame(content_frame, style='Surface.TFrame')
        button_frame.pack(pady=20)
        
        ttk.Button(button_frame, text='Сохранить', command=save_manufacture, style='Primary.TButton').pack(side='left', padx=10)
        ttk.Button(button_frame, text='Отмена', command=dialog.destroy, style='Secondary.TButton').pack(side='left', padx=10)
    
    def delete_manufacture(self):
        """Удаляет производителя"""
        selected = self.manufactures_tree.selection()
        if not selected:
            messagebox.showwarning('Внимание!', 'Выберите производителя')
            return
        
        item = selected[0]
        manufacture_id = self.manufactures_tree.item(item)['values'][0]
        manufacture_name = self.manufactures_tree.item(item)['values'][1]
        
        if messagebox.askyesno('Подтверждение', f'Удалить производителя "{manufacture_name}"?'):
            data = self.make_api_request('/manufactures/del_manufacture_by_id/', method='DELETE',
                                       params={'manufacture_id': manufacture_id})
            if data:
                messagebox.showinfo('Успех!', 'Производитель удален')
                self.load_manufactures()
    
    def init_components_tab(self):
        """Инициализирует вкладку компонентов"""
        main_frame = ttk.Frame(self.tab_components, style='Surface.TFrame')
        main_frame.pack(fill='both', expand=True)
        
        ttk.Label(main_frame, text='Управление компонентами', style='Header.TLabel').pack(anchor='w', pady=(0, 20))

        filter_frame = ttk.Frame(main_frame, style='Surface.TFrame')
        filter_frame.pack(fill='x', pady=(0, 20))

        ttk.Label(filter_frame, text='Тип:', style='Normal.TLabel').pack(side='left', padx=5)
        self.admin_type_filter = tk.StringVar()
        self.admin_type_combo = ttk.Combobox(filter_frame, textvariable=self.admin_type_filter, state='readonly', width=15, style='Modern.TCombobox')
        self.admin_type_combo.pack(side='left', padx=5)

        ttk.Label(filter_frame, text='Цена от:', style='Normal.TLabel').pack(side='left', padx=5)
        self.admin_min_price_filter = tk.StringVar()
        admin_min_price_entry = ttk.Entry(filter_frame, textvariable=self.admin_min_price_filter, width=8, style='Modern.TEntry')
        admin_min_price_entry.pack(side='left', padx=2)
        
        ttk.Label(filter_frame, text='до:', style='Normal.TLabel').pack(side='left', padx=2)
        self.admin_max_price_filter = tk.StringVar()
        admin_max_price_entry = ttk.Entry(filter_frame, textvariable=self.admin_max_price_filter, width=8, style='Modern.TEntry')
        admin_max_price_entry.pack(side='left', padx=5)
 
        ttk.Button(filter_frame, text='Применить', 
                command=self.admin_apply_filters, style='Primary.TButton').pack(side='left', padx=5)
        ttk.Button(filter_frame, text='Сбросить', 
                command=self.admin_reset_filters, style='Secondary.TButton').pack(side='left', padx=5)
     
        search_frame = ttk.Frame(main_frame, style='Surface.TFrame')
        search_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Label(search_frame, text='Поиск:', style='Normal.TLabel').pack(side='left', padx=5)
        self.admin_search_var = tk.StringVar()
        admin_search_entry = ttk.Entry(search_frame, textvariable=self.admin_search_var, width=40, style='Modern.TEntry')
        admin_search_entry.pack(side='left', padx=5)
        
        ttk.Button(search_frame, text='Найти', 
                command=self.admin_search_components, style='Primary.TButton').pack(side='left', padx=5)
        ttk.Button(search_frame, text='Сбросить', 
                command=self.admin_clear_search, style='Secondary.TButton').pack(side='left', padx=5)

        button_frame = ttk.Frame(main_frame, style='Surface.TFrame')
        button_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Button(button_frame, text='Добавить компонент', 
                  command=self.add_component, style='Primary.TButton').pack(side='left', padx=5)
        ttk.Button(button_frame, text='Обновить', 
                  command=self.load_components, style='Secondary.TButton').pack(side='left', padx=5)
        ttk.Button(button_frame, text='Редактировать', 
                  command=self.edit_component, style='Secondary.TButton').pack(side='left', padx=5)
        ttk.Button(button_frame, text='Удалить', 
                  command=self.delete_component, style='Secondary.TButton').pack(side='left', padx=5)

        columns = ('id', 'name', 'type', 'manufacture', 'price', 'stock')
        self.components_tree = ttk.Treeview(main_frame, columns=columns, show='headings', height=20)
        
        self.components_tree.heading('id', text='ID')
        self.components_tree.heading('name', text='Название')
        self.components_tree.heading('type', text='Тип')
        self.components_tree.heading('manufacture', text='Производитель')
        self.components_tree.heading('price', text='Цена')
        self.components_tree.heading('stock', text='Наличие')
        
        self.components_tree.column('id', width=50)
        self.components_tree.column('name', width=150)
        self.components_tree.column('type', width=100)
        self.components_tree.column('manufacture', width=120)
        self.components_tree.column('price', width=80)
        self.components_tree.column('stock', width=80)
        
        scrollbar = ttk.Scrollbar(main_frame, orient=tk.VERTICAL, command=self.components_tree.yview)
        self.components_tree.configure(yscroll=scrollbar.set)
        
        self.components_tree.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')

        admin_search_entry.bind('<Return>', lambda event: self.admin_search_components())

        self.admin_all_components = []
        self.admin_filtered_components = []
        self.admin_available_types = set()
        
        self.load_components()
    
    def load_components(self):
        """Загружает список компонентов"""
        data = self.make_api_request('/components/get_all/')
        if data:
            for item in self.components_tree.get_children():
                self.components_tree.delete(item)
            
            self.admin_all_components = data
            self.admin_filtered_components = data.copy()

            self.admin_available_types = set()
            for component in data:
                if component.get('type_name'):
                    self.admin_available_types.add(component['type_name'])

            self.admin_type_combo['values'] = ["Все типы"] + sorted(list(self.admin_available_types))
            self.admin_type_combo.set("Все типы")

            self.admin_display_filtered_components()
    
    def admin_apply_filters(self):
        """Применяет фильтры к списку компонентов"""
        if not hasattr(self, 'admin_all_components') or not self.admin_all_components:
            return
        
        filtered = self.admin_all_components.copy()

        selected_type = self.admin_type_filter.get()
        if selected_type and selected_type != "Все типы":
            filtered = [comp for comp in filtered if comp.get('type_name') == selected_type]

        min_price = self.admin_min_price_filter.get()
        if min_price:
            try:
                min_price_val = float(min_price)
                filtered = [comp for comp in filtered if comp.get('price', 0) >= min_price_val]
            except ValueError:
                pass

        max_price = self.admin_max_price_filter.get()
        if max_price:
            try:
                max_price_val = float(max_price)
                filtered = [comp for comp in filtered if comp.get('price', 0) <= max_price_val]
            except ValueError:
                pass

        search_term = self.admin_search_var.get().strip().lower()
        if search_term:
            filtered = [comp for comp in filtered if (
                search_term in comp['name'].lower() or 
                search_term in comp['type_name'].lower() or 
                search_term in comp['manufacture_name'].lower()
            )]
        
        self.admin_filtered_components = filtered
        self.admin_display_filtered_components()

    def admin_reset_filters(self):
        """Сбрасывает фильтры компонентов"""
        self.admin_type_filter.set('Все типы')
        self.admin_min_price_filter.set('')
        self.admin_max_price_filter.set('')
        self.admin_search_var.set('')
        self.admin_filtered_components = self.admin_all_components.copy()
        self.admin_display_filtered_components()

    def admin_display_filtered_components(self):
        """Отображает отфильтрованные компоненты"""
        for item in self.components_tree.get_children():
            self.components_tree.delete(item)
        
        for component in self.admin_filtered_components:
            self.components_tree.insert('', 'end', values=(
                component['id'],
                component['name'],
                component['type_name'],
                component['manufacture_name'],
                component['price'],
                component['stock_quantity']
            ))

    def admin_search_components(self):
        """Выполняет поиск компонентов"""
        self.admin_apply_filters()
        
        if self.admin_search_var.get().strip():
            messagebox.showinfo('Поиск', f'Найдено компонентов: {len(self.admin_filtered_components)}')

    def admin_clear_search(self):
        """Очищает поисковый запрос"""
        self.admin_search_var.set('')
        self.admin_apply_filters()
    
    def add_component(self):
        """Добавляет компонент"""
        dialog = tk.Toplevel(self.root)
        dialog.title('Добавление компонента')
        dialog.geometry('400x400')
        dialog.configure(bg=self.style.background_color)
        dialog.transient(self.root)
        dialog.grab_set()

        content_frame = ttk.Frame(dialog, style='Surface.TFrame', padding=20)
        content_frame.pack(fill='both', expand=True)

        ttk.Label(content_frame, text='Название:', style='Normal.TLabel').grid(row=0, column=0, padx=10, pady=5, sticky='e')
        name_entry = ttk.Entry(content_frame, width=30, style='Modern.TEntry')
        name_entry.grid(row=0, column=1, padx=10, pady=5)
        
        ttk.Label(content_frame, text='Тип:', style='Normal.TLabel').grid(row=1, column=0, padx=10, pady=5, sticky='e')
        type_entry = ttk.Entry(content_frame, width=30, style='Modern.TEntry')
        type_entry.grid(row=1, column=1, padx=10, pady=5)
        
        ttk.Label(content_frame, text='Производитель:', style='Normal.TLabel').grid(row=2, column=0, padx=10, pady=5, sticky='e')
        manufacture_entry = ttk.Entry(content_frame, width=30, style='Modern.TEntry')
        manufacture_entry.grid(row=2, column=1, padx=10, pady=5)
        
        ttk.Label(content_frame, text='Цена:', style='Normal.TLabel').grid(row=3, column=0, padx=10, pady=5, sticky='e')
        price_entry = ttk.Entry(content_frame, width=30, style='Modern.TEntry')
        price_entry.grid(row=3, column=1, padx=10, pady=5)
        
        ttk.Label(content_frame, text='Количество:', style='Normal.TLabel').grid(row=4, column=0, padx=10, pady=5, sticky='e')
        stock_entry = ttk.Entry(content_frame, width=30, style='Modern.TEntry')
        stock_entry.grid(row=4, column=1, padx=10, pady=5)
        
        def save_component():
            name = name_entry.get().strip()
            type_name = type_entry.get().strip()
            manufacture_name = manufacture_entry.get().strip()
            
            if not name:
                messagebox.showerror('Ошибка!', 'Введите название')
                return
            
            try:
                price = float(price_entry.get())
                stock = int(stock_entry.get())
            except ValueError:
                messagebox.showerror('Ошибка!', 'Цена и количество должны быть числами')
                return
            
            component_data = {
                'name': name,
                'type_name': type_name if type_name else None,
                'manufacture_name': manufacture_name if manufacture_name else None,
                'price': price,
                'stock_quantity': stock,
                'specification': []
            }
            
            data = self.make_api_request('/components/create/', method='POST', json_data=component_data)
            if data:
                messagebox.showinfo('Успех!', 'Компонент добавлен')
                dialog.destroy()
                self.load_components()
        
        button_frame = ttk.Frame(content_frame, style='Surface.TFrame')
        button_frame.grid(row=5, column=0, columnspan=2, pady=20)
        
        ttk.Button(button_frame, text='Сохранить', command=save_component, style='Primary.TButton').pack(side='left', padx=10)
        ttk.Button(button_frame, text='Отмена', command=dialog.destroy, style='Secondary.TButton').pack(side='left', padx=10)
    
    def edit_component(self):
        """Редактирует компонент"""
        selected = self.components_tree.selection()
        if not selected:
            messagebox.showwarning('Внимание!', 'Выберите компонент')
            return

        item = selected[0]
        component_id = self.components_tree.item(item)['values'][0]
        current_data = None
        for component in self.components_tree.get_children():
            if self.components_tree.item(component)['values'][0] == component_id:
                current_data = self.components_tree.item(component)['values']
                break
        
        if not current_data:
            messagebox.showerror('Ошибка!', 'Не удалось получить данные компонента')
            return

        dialog = tk.Toplevel(self.root)
        dialog.title('Редактирование компонента')
        dialog.geometry('500x600')
        dialog.configure(bg=self.style.background_color)
        dialog.transient(self.root)
        dialog.grab_set()

        content_frame = ttk.Frame(dialog, style='Surface.TFrame', padding=20)
        content_frame.pack(fill='both', expand=True)

        ttk.Label(content_frame, text='Название:', style='Normal.TLabel').grid(row=0, column=0, padx=10, pady=5, sticky='e')
        name_entry = ttk.Entry(content_frame, width=30, style='Modern.TEntry')
        name_entry.insert(0, current_data[1])
        name_entry.grid(row=0, column=1, padx=10, pady=5)

        ttk.Label(content_frame, text='Тип:', style='Normal.TLabel').grid(row=1, column=0, padx=10, pady=5, sticky='e')
        type_entry = ttk.Entry(content_frame, width=30, style='Modern.TEntry')
        type_entry.insert(0, current_data[2] if current_data[2] else '')
        type_entry.grid(row=1, column=1, padx=10, pady=5)

        ttk.Label(content_frame, text='Производитель:', style='Normal.TLabel').grid(row=2, column=0, padx=10, pady=5, sticky='e')
        manufacture_entry = ttk.Entry(content_frame, width=30, style='Modern.TEntry')
        manufacture_entry.insert(0, current_data[3] if current_data[3] else '')
        manufacture_entry.grid(row=2, column=1, padx=10, pady=5)

        ttk.Label(content_frame, text='Цена:', style='Normal.TLabel').grid(row=3, column=0, padx=10, pady=5, sticky='e')
        price_entry = ttk.Entry(content_frame, width=30, style='Modern.TEntry')
        price_entry.insert(0, str(current_data[4]))
        price_entry.grid(row=3, column=1, padx=10, pady=5)

        ttk.Label(content_frame, text='Количество на складе:', style='Normal.TLabel').grid(row=4, column=0, padx=10, pady=5, sticky='e')
        stock_entry = ttk.Entry(content_frame, width=30, style='Modern.TEntry')
        stock_entry.insert(0, str(current_data[5]))
        stock_entry.grid(row=4, column=1, padx=10, pady=5)

        spec_frame = ttk.LabelFrame(content_frame, text='Спецификация (JSON)')
        spec_frame.grid(row=5, column=0, columnspan=2, padx=10, pady=5, sticky='nsew')

        ttk.Label(spec_frame, text='Формат: [{"key": "название", "value": "значение"}, ...] или {"ключ": "значение", ...}', 
                font=('Arial', 8), foreground='gray').pack(anchor='w', padx=5, pady=2)
        
        spec_text = scrolledtext.ScrolledText(spec_frame, width=60, height=8, bg=self.style.surface_color, fg=self.style.text_primary)
        spec_text.pack(fill='both', expand=True, padx=5, pady=5)

        def load_component_details():
            """Загрузка данных о компоненте"""
            try:
                headers = {'token': self.token}
                response = requests.get(
                    f'{self.base_url}/components/get_by_id/',
                    headers=headers,
                    params={'component_id': component_id}
                )
                
                if response.status_code == 200:
                    component_data = response.json()
                    specification = component_data.get('specification', [])
                    if specification:
                        formatted_spec = json.dumps(specification, indent=2, ensure_ascii=False)
                        spec_text.insert('1.0', formatted_spec)
                else:
                    messagebox.showerror('Ошибка!', 'Не удалось загрузить детали компонента')
            except requests.exceptions.RequestException as e:
                messagebox.showerror('Ошибка!', f'Ошибка соединения: {e}')

        load_component_details()

        def validate_json(data):
            """Обработка json данных"""
            try:
                parsed = json.loads(data)
                if not isinstance(parsed, (dict, list)):
                    return False, "JSON должен быть объектом или массивом"
                return True, parsed
            except json.JSONDecodeError as e:
                return False, f"Ошибка JSON: {str(e)}"

        def save_component():
            """Сохранение изменений компонента"""
            name = name_entry.get().strip()
            type_name = type_entry.get().strip()
            manufacture_name = manufacture_entry.get().strip()
            
            if not name:
                messagebox.showerror('Ошибка!', 'Введите название компонента')
                return
            
            try:
                price = float(price_entry.get())
                stock_quantity = int(stock_entry.get())
            except ValueError:
                messagebox.showerror('Ошибка!', 'Цена и количество должны быть числами')
                return

            specification_text = spec_text.get('1.0', tk.END).strip()
            specification = []
            
            if specification_text:
                is_valid, result = validate_json(specification_text)
                if not is_valid:
                    messagebox.showerror('Ошибка!', result)
                    return
                specification = result

            component_data = {
                'new_name': name,
                'type_name': type_name if type_name else None,
                'manufacture_name': manufacture_name if manufacture_name else None,
                'price': price,
                'stock_quantity': stock_quantity,
                'specification': specification
            }

            data = self.make_api_request('/components/edit_by_id/', method='PUT',
                                    json_data=component_data,
                                    params={'component_id': component_id})
            if data:
                messagebox.showinfo('Успех!', 'Компонент обновлен')
                dialog.destroy()
                self.load_components()

        button_frame = ttk.Frame(content_frame, style='Surface.TFrame')
        button_frame.grid(row=6, column=0, columnspan=2, pady=20)

        ttk.Button(button_frame, text='Сохранить', command=save_component, 
                style='Primary.TButton').pack(side='left', padx=10)
        ttk.Button(button_frame, text='Отмена', command=dialog.destroy, 
                style='Secondary.TButton').pack(side='left', padx=10)

        dialog.grid_rowconfigure(5, weight=1)
        dialog.grid_columnconfigure(1, weight=1)

        help_label = ttk.Label(content_frame, 
                            text='Подсказка: Используйте валидный JSON формат. Пример: [{"Сокет": "LGA 1700"}, {"Чипсет": "Intel Z690"}]',
                            font=('Arial', 8),
                            foreground='blue')
        help_label.grid(row=7, column=0, columnspan=2, pady=5)
    
    def delete_component(self):
        """Удаляет компонент"""
        selected = self.components_tree.selection()
        if not selected:
            messagebox.showwarning('Внимание!', 'Выберите компонент')
            return
        
        item = selected[0]
        component_id = self.components_tree.item(item)['values'][0]
        component_name = self.components_tree.item(item)['values'][1]
        
        if messagebox.askyesno('Подтверждение', f'Удалить компонент "{component_name}"?'):
            data = self.make_api_request('/components/delete_by_id/', method='DELETE',
                                       params={'component_id': component_id})
            if data:
                messagebox.showinfo('Успех!', 'Компонент удален')
                self.load_components()
    
    def init_configurations_tab(self):
        """Инициализирует вкладку конфигураций"""
        main_frame = ttk.Frame(self.tab_configurations, style='Surface.TFrame')
        main_frame.pack(fill='both', expand=True)
        
        ttk.Label(main_frame, text='Управление конфигурациями', style='Header.TLabel').pack(anchor='w', pady=(0, 20))

        search_frame = ttk.Frame(main_frame, style='Surface.TFrame')
        search_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Label(search_frame, text='Поиск по названию:', style='Normal.TLabel').pack(side='left', padx=5)
        
        self.config_search_var = tk.StringVar()
        search_entry = ttk.Entry(search_frame, textvariable=self.config_search_var, width=30, style='Modern.TEntry')
        search_entry.pack(side='left', padx=5)
        
        ttk.Button(search_frame, text='Найти', 
                  command=self.search_configurations, style='Primary.TButton').pack(side='left', padx=5)
        ttk.Button(search_frame, text='Обновить', 
                  command=self.load_configurations, style='Secondary.TButton').pack(side='left', padx=5)

        button_frame = ttk.Frame(main_frame, style='Surface.TFrame')
        button_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Button(button_frame, text='Просмотреть', 
                  command=self.view_configuration, style='Primary.TButton').pack(side='left', padx=5)
        ttk.Button(button_frame, text='Удалить', 
                  command=self.delete_configuration, style='Secondary.TButton').pack(side='left', padx=5)

        columns = ('id', 'user', 'name', 'description', 'created')
        self.configurations_tree = ttk.Treeview(main_frame, columns=columns, show='headings', height=20)
        
        self.configurations_tree.heading('id', text='ID')
        self.configurations_tree.heading('user', text='Пользователь')
        self.configurations_tree.heading('name', text='Название')
        self.configurations_tree.heading('description', text='Описание')
        self.configurations_tree.heading('created', text='Дата создания')
        
        self.configurations_tree.column('id', width=50)
        self.configurations_tree.column('user', width=120)
        self.configurations_tree.column('name', width=150)
        self.configurations_tree.column('description', width=200)
        self.configurations_tree.column('created', width=120)
        
        scrollbar = ttk.Scrollbar(main_frame, orient=tk.VERTICAL, command=self.configurations_tree.yview)
        self.configurations_tree.configure(yscroll=scrollbar.set)
        
        self.configurations_tree.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')

        search_entry.bind('<Return>', lambda event: self.search_configurations())
        
        self.all_configurations = []
        self.load_configurations()

    def search_configurations(self):
        """Выполняет поиск конфигураций"""
        search_text = self.config_search_var.get().strip().lower()
        
        if not search_text:
            self.display_configurations(self.all_configurations)
            return

        filtered_configs = [
            config for config in self.all_configurations 
            if search_text in config['name_config'].lower()
        ]
        
        self.display_configurations(filtered_configs)
        
        messagebox.showinfo('Поиск', f'Найдено конфигураций: {len(filtered_configs)}')

    def load_configurations(self):
        """Загружает список конфигураций"""
        data = self.make_api_request('/configurations/admin/get_all/')
        if data:
            self.all_configurations = data
            self.display_configurations(data)
            
    def display_configurations(self, configurations):
        """Отображает список конфигураций"""
        for item in self.configurations_tree.get_children():
            self.configurations_tree.delete(item)
            
        for config in configurations:
            created_at = config.get('created_at', '')
            if created_at:
                try:
                    created_at = created_at.split('T')[0]
                except:
                    pass
            
            self.configurations_tree.insert('', 'end', values=(
                config['id'],
                config['user_name'],
                config['name_config'],
                config['description'] or '-',
                created_at
            ))
    
    def view_configuration(self):
        """Просматривает детали конфигурации"""
        selected = self.configurations_tree.selection()
        if not selected:
            messagebox.showwarning('Внимание!', 'Выберите конфигурацию')
            return
        
        item = selected[0]
        config_id = self.configurations_tree.item(item)['values'][0]
        
        data = self.make_api_request(f'/configurations/admin/{config_id}/components/')
        if data:
            dialog = tk.Toplevel(self.root)
            dialog.title(f'Компоненты конфигурации #{config_id}')
            dialog.geometry('1000x600')
            dialog.configure(bg=self.style.background_color)
            dialog.transient(self.root)
            dialog.grab_set()
            
            content_frame = ttk.Frame(dialog, style='Surface.TFrame', padding=20)
            content_frame.pack(fill='both', expand=True)
            
            columns = ('component', 'type', 'manufacture', 'price', 'quantity', 'total')
            tree = ttk.Treeview(content_frame, columns=columns, show='headings', height=15)
            
            tree.heading('component', text='Компонент')
            tree.heading('type', text='Тип')
            tree.heading('manufacture', text='Производитель')
            tree.heading('price', text='Цена')
            tree.heading('quantity', text='Количество')
            tree.heading('total', text='Итого')
            
            tree.column('component', width=150)
            tree.column('type', width=100)
            tree.column('manufacture', width=120)
            tree.column('price', width=80)
            tree.column('quantity', width=80)
            tree.column('total', width=80)
            
            scrollbar = ttk.Scrollbar(content_frame, orient=tk.VERTICAL, command=tree.yview)
            tree.configure(yscroll=scrollbar.set)
            
            tree.pack(side='left', fill='both', expand=True, padx=10, pady=10)
            scrollbar.pack(side='right', fill='y')
            
            total_amount = 0
            for component in data:
                total_price = component['total_price']
                total_amount += total_price
                
                tree.insert('', 'end', values=(
                    component['component_name'],
                    component['type_name'] or '-',
                    component['manufacture_name'] or '-',
                    f"{component['price']:.2f} руб.",
                    component['quantity'],
                    f"{total_price:.2f} руб."
                ))
            
            ttk.Label(content_frame, text=f'Общая сумма: {total_amount:.2f} руб.', style='Header.TLabel').pack(pady=10)
    
    def delete_configuration(self):
        """Удаляет конфигурацию"""
        selected = self.configurations_tree.selection()
        if not selected:
            messagebox.showwarning('Внимание!', 'Выберите конфигурацию')
            return
        
        item = selected[0]
        config_id = self.configurations_tree.item(item)['values'][0]
        config_name = self.configurations_tree.item(item)['values'][2]
        
        if messagebox.askyesno('Подтверждение', f'Удалить конфигурацию "{config_name}"?'):
            data = self.make_api_request('/configurations/admin/delete_by_id/', method='DELETE',
                                       params={'config_id': config_id})
            if data:
                messagebox.showinfo('Успех!', 'Конфигурация удалена')
                self.load_configurations()
    
    def init_orders_tab(self):
        """Инициализирует вкладку заказов"""
        main_frame = ttk.Frame(self.tab_orders, style='Surface.TFrame')
        main_frame.pack(fill='both', expand=True)
        
        ttk.Label(main_frame, text='Управление заказами', style='Header.TLabel').pack(anchor='w', pady=(0, 20))

        search_frame = ttk.Frame(main_frame, style='Surface.TFrame')
        search_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Label(search_frame, text='Поиск по email:', style='Normal.TLabel').pack(side='left', padx=5)
        
        self.order_search_var = tk.StringVar()
        search_entry = ttk.Entry(search_frame, textvariable=self.order_search_var, width=30, style='Modern.TEntry')
        search_entry.pack(side='left', padx=5)
        
        ttk.Button(search_frame, text='Найти', 
                  command=self.search_orders, style='Primary.TButton').pack(side='left', padx=5)
        ttk.Button(search_frame, text='Обновить', 
                  command=self.load_orders, style='Secondary.TButton').pack(side='left', padx=5)

        button_frame = ttk.Frame(main_frame, style='Surface.TFrame')
        button_frame.pack(fill='x', pady=(0, 20))
        
        ttk.Button(button_frame, text='Изменить статус', 
                  command=self.change_order_status, style='Primary.TButton').pack(side='left', padx=5)
        ttk.Button(button_frame, text='Удалить', 
                  command=self.delete_order, style='Secondary.TButton').pack(side='left', padx=5)

        columns = ('id', 'user', 'date', 'total', 'status')
        self.orders_tree = ttk.Treeview(main_frame, columns=columns, show='headings', height=20)
        
        self.orders_tree.heading('id', text='ID заказа')
        self.orders_tree.heading('user', text='Пользователь')
        self.orders_tree.heading('date', text='Дата')
        self.orders_tree.heading('total', text='Сумма')
        self.orders_tree.heading('status', text='Статус')
        
        self.orders_tree.column('id', width=80)
        self.orders_tree.column('user', width=120)
        self.orders_tree.column('date', width=100)
        self.orders_tree.column('total', width=100)
        self.orders_tree.column('status', width=120)
        
        scrollbar = ttk.Scrollbar(main_frame, orient=tk.VERTICAL, command=self.orders_tree.yview)
        self.orders_tree.configure(yscroll=scrollbar.set)
        
        self.orders_tree.pack(side='left', fill='both', expand=True)
        scrollbar.pack(side='right', fill='y')

        search_entry.bind('<Return>', lambda event: self.search_orders())
        
        self.all_orders = []
        self.load_orders()

    def search_orders(self):
        """Выполняет поиск заказов"""
        search_text = self.order_search_var.get().strip().lower()
        
        if not search_text:
            self.display_orders(self.all_orders)
            return

        filtered_orders = [
            order for order in self.all_orders 
            if search_text in order['user_login'].lower()
        ]
        
        self.display_orders(filtered_orders)
        
        messagebox.showinfo('Поиск', f'Найдено заказов: {len(filtered_orders)}')

    def load_orders(self):
        """Загружает список заказов"""
        data = self.make_api_request('/orders/admin/get_all')
        if data:
            self.all_orders = data
            self.display_orders(data)
            
    def display_orders(self, orders):
        """Отображает список заказов"""
        for item in self.orders_tree.get_children():
            self.orders_tree.delete(item)
            
        for order in orders:
            order_date = order.get('order_date', '')
            if order_date:
                try:
                    order_date = order_date.split('T')[0]
                except:
                    pass
            
            total_amount = order.get('total_amount', 0)
            total_text = f"{total_amount:.2f} руб." if total_amount else "0 руб."
            
            self.orders_tree.insert('', 'end', values=(
                order['id'],
                order['user_login'],
                order_date,
                total_text,
                order.get('status_name', 'Неизвестно')
            ))
    
    def change_order_status(self):
        """Изменяет статус заказа"""
        selected = self.orders_tree.selection()
        if not selected:
            messagebox.showwarning('Внимание!', 'Выберите заказ')
            return
        
        item = selected[0]
        order_id = self.orders_tree.item(item)['values'][0]
        current_status = self.orders_tree.item(item)['values'][4]

        statuses_data = self.make_api_request('/order_status/get_all/')
        if not statuses_data:
            return
        
        status_names = [status['name'] for status in statuses_data]
        
        dialog = tk.Toplevel(self.root)
        dialog.title('Изменение статуса заказа')
        dialog.geometry('400x200')
        dialog.configure(bg=self.style.background_color)
        dialog.transient(self.root)
        dialog.grab_set()
        
        content_frame = ttk.Frame(dialog, style='Surface.TFrame', padding=20)
        content_frame.pack(fill='both', expand=True)
        
        ttk.Label(content_frame, text=f'Заказ: #{order_id}', style='Normal.TLabel').pack(pady=10)
        ttk.Label(content_frame, text=f'Текущий статус: {current_status}', style='Normal.TLabel').pack(pady=5)
        
        status_var = tk.StringVar(value=current_status)
        status_combo = ttk.Combobox(content_frame, textvariable=status_var, state='readonly', style='Modern.TCombobox')
        status_combo['values'] = status_names
        status_combo.pack(pady=10)
        
        def save_status():
            new_status = status_var.get()
            if new_status == current_status:
                messagebox.showinfo('Информация', 'Статус не изменен')
                dialog.destroy()
                return

            status_id = None
            for status in statuses_data:
                if status['name'] == new_status:
                    status_id = status['id']
                    break
            
            if status_id is None:
                messagebox.showerror('Ошибка!', 'Статус не найден')
                return
            
            data = self.make_api_request('/orders/admin/edit_order_status/', method='PUT',
                                       json_data={'status_id': status_id},
                                       params={'order_id': order_id})
            if data:
                messagebox.showinfo('Успех!', 'Статус заказа изменен')
                dialog.destroy()
                self.load_orders()
        
        ttk.Button(content_frame, text='Сохранить', command=save_status, style='Primary.TButton').pack(pady=10)
    
    def delete_order(self):
        """Удаляет заказ"""
        selected = self.orders_tree.selection()
        if not selected:
            messagebox.showwarning('Внимание!', 'Выберите заказ')
            return
        
        item = selected[0]
        order_id = self.orders_tree.item(item)['values'][0]
        
        if messagebox.askyesno('Подтверждение', f'Удалить заказ #{order_id}?'):
            data = self.make_api_request('/orders/admin/delete_order/', method='DELETE',
                                       params={'order_id': order_id})
            if data:
                messagebox.showinfo('Успех!', 'Заказ удален')
                self.load_orders()

if __name__ == '__main__':
    root = tk.Tk()
    app = ModernAuthApp(root)
    root.mainloop()
